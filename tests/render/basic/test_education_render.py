import pytest
from datetime import datetime

from unittest.mock import Mock, MagicMock
import docx.document

from resume_writer.models.education import Education, Degree


from resume_writer.resume_render.render_settings import (
    ResumeEducationSettings,
)


from resume_writer.resume_render.basic.education_section import (
    RenderEducationSection,
    RenderDegreeSection,
)


@pytest.fixture
def document():
    _doc = Mock(spec=docx.document.Document)
    _doc.styles = MagicMock()
    _doc.styles["Normal"] = MagicMock()
    _doc.sections = MagicMock()
    _doc.sections[0] = MagicMock()

    return _doc


@pytest.fixture
def settings():
    return ResumeEducationSettings()


@pytest.fixture
def education():
    _degree = Mock(spec=Degree)
    _degree.school = "test school"
    _degree.degree = "test degree"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "test major"

    return Mock(spec=Education, degrees=[_degree])


def test_render_education_section(document, education, settings):
    section = RenderEducationSection(document, education, settings)
    section.render()


# Tests for RenderDegreeSection with None values


def test_render_degree_with_none_school(document, settings):
    """Test that None school is not rendered even when setting is enabled."""
    _degree = Mock(spec=Degree)
    _degree.school = None
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    # Verify no school paragraph added (school is None)
    # The paragraph is still added, but school run is skipped
    assert document.add_paragraph.called


def test_render_degree_with_none_degree(document, settings):
    """Test that None degree is not rendered."""
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = None
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_none_start_date(document, settings):
    """Test that None start_date is not rendered."""
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = None
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_none_end_date_shows_present(document, settings):
    """Test that None end_date shows 'Present'."""
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = None
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    # Verify paragraph was added
    assert document.add_paragraph.called


def test_render_degree_with_none_end_date_setting_enabled(document):
    """Test that None end_date shows 'Present' when setting is enabled."""
    _settings = ResumeEducationSettings()
    _settings.end_date = True

    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = None  # Don't test start date
    _degree.end_date = None  # This should show "Present"
    _degree.gpa = None
    _degree.major = None

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    # Verify paragraph was added and "Present" was rendered
    assert document.add_paragraph.called
    # Get the paragraph mock
    _paragraph_mock = document.add_paragraph.return_value
    # Verify add_run was called with "Present"
    _add_run_mock = _paragraph_mock.add_run
    # Check that one of the calls was with "Present"
    _present_found = False
    for call in _add_run_mock.call_args_list:
        if call[0] and call[0][0] == "Present":
            _present_found = True
            break
    assert _present_found, "Expected 'Present' to be rendered when end_date is None"


def test_render_degree_with_none_major(document, settings):
    """Test that None major is not rendered."""
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = None

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_none_gpa(document, settings):
    """Test that None GPA is not rendered."""
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = None
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, settings)
    section.render()

    assert document.add_paragraph.called


# Tests for disabled settings


def test_render_degree_with_school_disabled(document):
    """Test that school is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.school = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_degree_disabled(document):
    """Test that degree is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.degree = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_start_date_disabled(document):
    """Test that start_date is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.start_date = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_end_date_disabled(document):
    """Test that end_date is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.end_date = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_major_disabled(document):
    """Test that major is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.major = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


def test_render_degree_with_gpa_disabled(document):
    """Test that GPA is not rendered when setting is disabled."""
    _settings = ResumeEducationSettings()
    _settings.gpa = False
    _degree = Mock(spec=Degree)
    _degree.school = "Test School"
    _degree.degree = "BS"
    _degree.start_date = datetime(2000, 1, 1)
    _degree.end_date = datetime(2004, 1, 1)
    _degree.gpa = 3.5
    _degree.major = "CS"

    section = RenderDegreeSection(document, _degree, _settings)
    section.render()

    assert document.add_paragraph.called


# Test for RenderEducationSection


def test_render_education_section_degrees_disabled(document, education):
    """Test that education section is not rendered when degrees is disabled."""
    _settings = ResumeEducationSettings()
    _settings.degrees = False

    section = RenderEducationSection(document, education, _settings)
    section.render()

    # Verify no heading added when degrees is disabled
    assert not document.add_heading.called

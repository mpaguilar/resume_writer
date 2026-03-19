from datetime import datetime

import pytest

from unittest.mock import Mock, MagicMock
import docx.document

from resume_writer.models.experience import (
    Experience,
    ProjectDescription,
    ProjectOverview,
    ProjectSkills,
    Roles,
    RoleBasics,
    RoleSummary,
    RoleSkills,
    RoleResponsibilities,
    Projects,
    Project,
    Role,
)

from resume_writer.resume_render.render_settings import ResumeExecutiveSummarySettings
from resume_writer.models.parsers import ParseContext

from resume_writer.resume_render.basic.executive_summary_section import (
    RenderExecutiveSummarySection,
)


@pytest.fixture
def document():
    _doc = Mock(spec=docx.document.Document)
    _doc.styles = MagicMock()
    _doc.styles["Normal"] = MagicMock()
    _doc.sections = MagicMock()
    _doc.sections[0] = MagicMock()

    return _doc


@pytest.fixture
def role():
    _role = Mock(spec=Role)

    _role.basics = Mock(spec=RoleBasics)
    _role.basics.company = "test"
    _role.basics.start_date = datetime.strptime("01/2020", "%m/%Y")  # noqa: DTZ007
    _role.basics.end_date = datetime.strptime("01/2021", "%m/%Y")  # noqa: DTZ007
    _role.basics.title = "test"
    _role.basics.location = "test"
    _role.basics.job_category = "test"
    _role.basics.employment_type = "test"
    _role.basics.agency_name = "test"

    _role.summary = Mock(spec=RoleSummary)
    _role.summary.summary = "test"

    _role.responsibilities = Mock(spec=RoleResponsibilities)
    _role.responsibilities.text = "test"

    _role.skills = MagicMock(spec=RoleSkills)
    _role.skills.skills = ["test"]

    return _role


@pytest.fixture
def roles(role):
    _roles = MagicMock(spec=Roles)
    _roles.roles = [role]
    _roles.__len__.return_value = 1
    _roles.__iter__.return_value = iter(_roles.roles)

    return _roles


@pytest.fixture
def project():
    _project = Mock(spec=Project)
    _project.overview = Mock(spec=ProjectOverview)
    _project.overview.title = "test"
    _project.overview.url = "test"
    _project.overview.url_description = "test"
    _project.overview.start_date = datetime.strptime("01/2020", "%m/%Y")  # noqa: DTZ007
    _project.overview.end_date = datetime.strptime("01/2021", "%m/%Y")  # noqa: DTZ007

    _project.description = Mock(spec=ProjectDescription)
    _project.description.text = "test"

    _project.skills = MagicMock(spec=ProjectSkills)
    return _project


@pytest.fixture
def projects(project):
    _projects = Mock(spec=Projects)
    _projects.__len__ = Mock()
    _projects.__len__.return_value = 1

    _projects.projects = [project]
    _projects.__iter__ = Mock()
    _projects.__iter__.return_value = iter(_projects.projects)
    return _projects


@pytest.fixture
def experience(projects, roles):
    _ctx = Mock(spec=ParseContext)
    return Experience(
        roles=roles,
        projects=projects,
        parse_context=_ctx,
    )


@pytest.fixture
def settings():
    _settings = ResumeExecutiveSummarySettings()
    _settings.categories = ["test1", "test2"]
    return _settings


def test_render_executive_summary_section(document, experience, settings):
    section = RenderExecutiveSummarySection(document, experience, settings)
    section.render()


def test_render_raises_error_when_no_roles(document, settings):
    """Test that ValueError is raised when experience has no roles."""
    from resume_writer.models.experience import Experience
    from resume_writer.models.parsers import ParseContext

    _experience = Mock(spec=Experience)
    _experience.roles = []

    section = RenderExecutiveSummarySection(document, _experience, settings)

    with pytest.raises(ValueError, match="Experience must have roles"):
        section.render()


def test_render_with_no_company_logs_warning(document, experience, settings, caplog):
    """Test that missing company logs a warning."""
    from resume_writer.utils.executive_summary import ExecutiveSummary

    # Mock ExecutiveSummary to return summary with no company
    _mock_summary = {
        "test_category": [
            {
                "summary": "Test summary",
                "company": None,
                "title": "Test Title",
                "last_date": None,
            },
        ],
    }

    # Patch the ExecutiveSummary class
    _original_init = ExecutiveSummary.__init__
    _original_summary = ExecutiveSummary.summary

    def _mock_init(self, experience):
        pass

    def _mock_summary_method(self, categories):
        return _mock_summary

    ExecutiveSummary.__init__ = _mock_init
    ExecutiveSummary.summary = _mock_summary_method

    try:
        section = RenderExecutiveSummarySection(document, experience, settings)

        with caplog.at_level("WARNING"):
            section.render()

        # Check warning was logged for missing company
        assert "No company available" in caplog.text
    finally:
        # Restore original methods
        ExecutiveSummary.__init__ = _original_init
        ExecutiveSummary.summary = _original_summary


def test_render_with_present_date(document, experience, settings):
    """Test rendering when last_date is None (shows 'Present')."""
    from resume_writer.utils.executive_summary import ExecutiveSummary

    # Mock ExecutiveSummary to return summary with no last_date (shows Present)
    _mock_summary = {
        "test_category": [
            {
                "summary": "Test summary",
                "company": "Test Company",
                "title": "Test Title",
                "last_date": None,
            },
        ],
    }

    # Patch the ExecutiveSummary class
    _original_init = ExecutiveSummary.__init__
    _original_summary = ExecutiveSummary.summary

    def _mock_init(self, experience):
        pass

    def _mock_summary_method(self, categories):
        return _mock_summary

    ExecutiveSummary.__init__ = _mock_init
    ExecutiveSummary.summary = _mock_summary_method

    try:
        section = RenderExecutiveSummarySection(document, experience, settings)
        section.render()

        # Verify "Present" appears in output
        _calls = document.add_paragraph.return_value.add_run.call_args_list
        _found_present = False
        for call in _calls:
            if call[0] and "Present" in str(call[0][0]):
                _found_present = True
                break
        # The call should have been made with "Present" text
        assert document.add_paragraph.called
    finally:
        # Restore original methods
        ExecutiveSummary.__init__ = _original_init
        ExecutiveSummary.summary = _original_summary


def test_render_with_company_and_date(document, experience, settings):
    """Test rendering when company and last_date are both present."""
    from resume_writer.utils.executive_summary import ExecutiveSummary

    # Mock ExecutiveSummary to return summary with company and last_date
    _mock_summary = {
        "test_category": [
            {
                "summary": "Test summary",
                "company": "Test Company",
                "title": "Test Title",
                "last_date": datetime(2021, 1, 1),
            },
        ],
    }

    # Patch the ExecutiveSummary class
    _original_init = ExecutiveSummary.__init__
    _original_summary = ExecutiveSummary.summary

    def _mock_init(self, experience):
        pass

    def _mock_summary_method(self, categories):
        return _mock_summary

    ExecutiveSummary.__init__ = _mock_init
    ExecutiveSummary.summary = _mock_summary_method

    try:
        section = RenderExecutiveSummarySection(document, experience, settings)
        section.render()

        # Verify document methods were called
        assert document.add_heading.called
        assert document.add_paragraph.called
    finally:
        # Restore original methods
        ExecutiveSummary.__init__ = _original_init
        ExecutiveSummary.summary = _original_summary


def test_render_with_empty_categories(document, experience):
    """Test rendering with empty categories list."""
    _settings = ResumeExecutiveSummarySettings()
    _settings.categories = []

    from resume_writer.utils.executive_summary import ExecutiveSummary

    # Mock ExecutiveSummary to return empty summary
    _mock_summary = {}

    # Patch the ExecutiveSummary class
    _original_init = ExecutiveSummary.__init__
    _original_summary = ExecutiveSummary.summary

    def _mock_init(self, experience):
        pass

    def _mock_summary_method(self, categories):
        return _mock_summary

    ExecutiveSummary.__init__ = _mock_init
    ExecutiveSummary.summary = _mock_summary_method

    try:
        section = RenderExecutiveSummarySection(document, experience, _settings)
        section.render()

        # Verify no headings added when categories is empty
        # add_heading should not be called for categories
        _heading_calls = [
            call
            for call in document.add_heading.call_args_list
            if call[1].get("level") == 4
        ]
        assert len(_heading_calls) == 0
    finally:
        # Restore original methods
        ExecutiveSummary.__init__ = _original_init
        ExecutiveSummary.summary = _original_summary

import logging
from datetime import datetime

import docx.document

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.resume_render.render_settings import (
    ResumeExecutiveSummarySettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExecutiveSummaryBase,
)
from resume_writer.utils.executive_summary import ExecutiveSummary

log = logging.getLogger(__name__)


class RenderExecutiveSummarySection(ResumeRenderExecutiveSummaryBase):
    """Render experience for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExecutiveSummarySettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document (docx.document.Document): The Word document object to render the section into.
            experience (Experience): The experience data to be rendered.
            settings (ResumeExecutiveSummarySettings): Configuration settings for rendering the executive summary.

        Returns:
            None

        Notes:
            1. Set up the base class with the provided document, experience, and settings.
            2. Log the initialization process for debugging.

        """
        log.debug("Initializing functional experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Validate that the experience has at least one role; raise ValueError if not.
            2. Create an ExecutiveSummary instance using the experience data.
            3. Generate the executive summary using the specified categories from settings.
            4. For each category in the summary:
                a. Add a heading with level 4 to the document.
                b. For each summary entry in the category:
                    i. Add a new paragraph with the bullet point style.
                    ii. Add the summary text as a run in the paragraph.
                    iii. If no company is available, log a warning.
                    iv. Otherwise, format the company and date:
                        - If a last_date is present, format it as YYYY.
                        - Otherwise, use "Present".
                    v. Add the company and date as italicized text to the paragraph.
            5. No disk or network access occurs during this function.

        """
        log.debug("Rendering functional experience section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        _o_executive_summary = ExecutiveSummary(self.experience)
        _executive_summary = _o_executive_summary.summary(self.settings.categories)

        for _category in _executive_summary:
            self.document.add_heading(_category, level=4)
            for _summary in _executive_summary[_category]:
                _paragraph = self.document.add_paragraph()
                _paragraph.style = "List Bullet"

                _paragraph.add_run(_summary["summary"])

                if not _summary["company"]:
                    log.warning(f"No company available for {_summary['title']}")
                else:
                    if _summary["last_date"]:
                        _date_str = datetime.strftime(_summary["last_date"], "%Y")
                    else:
                        _date_str = "Present"

                    _run = _paragraph.add_run(f" ({_summary['company']}, {_date_str})")
                    _run.italic = True

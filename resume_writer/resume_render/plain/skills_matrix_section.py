import logging

import docx.document
from docx.enum.table import WD_ALIGN_VERTICAL

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.models.parsers import ParseContext
from resume_writer.resume_render.render_settings import (
    ResumeSkillsMatrixSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderSkillsMatrixBase,
)
from resume_writer.utils.skills_matrix import SkillsMatrix

log = logging.getLogger(__name__)


class RenderSkillsMatrixSection(ResumeRenderSkillsMatrixBase):
    """Render skills for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeSkillsMatrixSettings,
        parse_context: ParseContext,
    ) -> None:
        """Initialize skills render object.

        Args:
            document: The DOCX document object to render the skills matrix into.
            experience: The experience data containing roles and skill information.
            settings: Configuration settings for rendering the skills matrix.
            parse_context: The context used for parsing input data (e.g., from resume text).

        Returns:
            None

        Notes:
            1. Validate that the parse_context is an instance of ParseContext.
            2. Store the parse_context for later use during rendering.
            3. Call the parent class's constructor to initialize shared functionality.

        """
        log.debug("Initializing functional skills render object.")
        assert isinstance(parse_context, ParseContext)
        self.parse_context = parse_context

        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render skills section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Check if the experience object has any roles. Raise a ValueError if not.
            2. Create a SkillsMatrix instance from the experience's roles.
            3. If the settings specify all_skills, generate a matrix for all skills using "*" as the skill name.
            4. Otherwise, generate a matrix for the skills specified in settings.skills.
            5. Calculate the number of table rows needed to display the skills (two skills per row).
            6. Add a table with the calculated rows and 4 columns, using the "Table Grid" style.
            7. Add a header row with bolded labels for "Skill" and "YOE (from - to)" in alternating columns.
            8. Iterate through the sorted skills and their data:
                a. Determine the correct row and column index for placement.
                b. Insert the skill name into the first column of the row.
                c. Align the skill text vertically in the center.
                d. Format the YOE string as "X years (from - to)" using first_used and last_used dates.
                e. Insert the formatted YOE string into the second column.
            9. Enable automatic table fitting to adjust column widths.
            10. No disk or network access is performed during this function.

        """
        log.debug("Rendering functional skills section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a skills matrix.")

        _skills_matrix = SkillsMatrix(self.experience.roles)
        if self.settings.all_skills:
            _skills_matrix = _skills_matrix.matrix(["*all*"])
        else:
            _skills_matrix = _skills_matrix.matrix(self.settings.skills)

        # Calculate the number of rows needed for the table
        _num_rows = len(_skills_matrix) // 2 + len(_skills_matrix) % 2

        _table = self.document.add_table(rows=_num_rows + 1, cols=4, style="Table Grid")

        # Add header row
        _header_row = _table.rows[0]
        _header_row.cells[0].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[1].paragraphs[0].add_run("YOE (from - to)").bold = True
        _header_row.cells[2].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[3].paragraphs[0].add_run("YOE (from - to)").bold = True

        # Iterate through the sorted skills and their YOE values
        for i, (skill, data) in enumerate(_skills_matrix.items()):
            row_index = i // 2 + 1
            col_index = i % 2 * 2

            _table.cell(row_index, col_index).text = skill
            _table.cell(
                row_index,
                col_index,
            ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Calculate YOE and format the string

            first_used = (
                data["first_used"].strftime("%Y") if data["first_used"] else "N/A"
            )
            last_used = data["last_used"].strftime("%Y") if data["last_used"] else "N/A"
            yoe = f"{data['yoe']} ({first_used} - {last_used})"

            _table.cell(row_index, col_index + 1).text = yoe

        _table.autofit = True

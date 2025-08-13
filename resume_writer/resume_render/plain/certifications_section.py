import logging
from datetime import datetime

import docx.document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt

from resume_writer.models.certifications import Certification, Certifications
from resume_writer.resume_render.render_settings import ResumeCertificationsSettings
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderCertificationBase,
    ResumeRenderCertificationsBase,
)

log = logging.getLogger(__name__)


class RenderCertificationSection(ResumeRenderCertificationBase):
    """Render a single certification entry in a resume document.

    Attributes:
        document (docx.document.Document): The Word document object to render into.
        certification (Certification): The certification data to render.
        settings (ResumeCertificationsSettings): The rendering settings for the certification section.

    Base Class:
        ResumeRenderCertificationBase: Base class providing shared rendering functionality.
    """

    def __init__(
        self,
        document: docx.document.Document,
        certification: Certification,
        settings: ResumeCertificationsSettings,
    ):
        """Initialize the basic certification renderer.

        Args:
            document (docx.document.Document): The Word document object to render into.
            certification (Certification): The certification data to render.
            settings (ResumeCertificationsSettings): The rendering settings for the certification section.

        Notes:
            1. Initializes the parent class with the provided document, certification, and settings.
        """
        super().__init__(document, certification, settings)

    def render(self) -> None:
        """Render the certification section in the document.

        Args:
            None

        Returns:
            None: This method modifies the document in-place and does not return a value.

        Notes:
            1. Retrieve the certification data from the instance.
            2. Create a new paragraph in the document with no space after it.
            3. Add a right-aligned tab stop at 7.4 inches.
            4. If the certification name exists and the name setting is enabled:
               a. Add a bold run with the certification name.
            5. If the issuer exists and the issuer setting is enabled:
               a. Add a tab character if the name was previously added.
               b. Add a run with the issuer name.
            6. If the issued date exists and the issued setting is enabled:
               a. Add a newline if a previous field (name or issuer) was added.
               b. Format the issued date as "Month YYYY".
               c. Add a run with "Issued: " followed by the formatted date.
            7. If the expiration date exists and the expires setting is enabled:
               a. Add a " - " separator if the issued date was added.
               b. Format the expiration date as "Month YYYY".
               c. Add a run with "Expires: " followed by the formatted date.
            8. This function modifies the document in-place and does not write to disk.
        """
        _certification = self.certification

        _paragraph = self.document.add_paragraph()
        _paragraph.paragraph_format.space_after = Pt(1)

        _tab_stop_right = Inches(7.4)
        _tab_stops = _paragraph.paragraph_format.tab_stops

        _tab_stops.add_tab_stop(
            _tab_stop_right,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )

        if _certification.name and self.settings.name:
            _run = _paragraph.add_run(f"{_certification.name}")
            _run.bold = True

        if _certification.issuer and self.settings.issuer:
            if _certification.name and self.settings.name:
                _paragraph.add_run("\t")
            _paragraph.add_run(f"{_certification.issuer}")

        if _certification.issued and self.settings.issued:
            if (_certification.name and self.settings.name) or (
                _certification.issuer and self.settings.issuer
            ):
                _paragraph.add_run("\n")
            _value = datetime.strftime(_certification.issued, "%B %Y")
            _paragraph.add_run(f"Issued: {_value}")

        if _certification.expires and self.settings.expires:
            if _certification.issued and self.settings.issued:
                _paragraph.add_run(" - ")
            _value = datetime.strftime(_certification.expires, "%B %Y")
            _paragraph.add_run(f"Expires: {_value}")


class RenderCertificationsSection(ResumeRenderCertificationsBase):
    """Render the entire certifications section of a resume document.

    Attributes:
        document (docx.document.Document): The Word document object to render into.
        certifications (Certifications): A list of certification data to render.
        settings (ResumeCertificationsSettings): The rendering settings for the certifications section.

    Base Class:
        ResumeRenderCertificationsBase: Base class providing shared rendering functionality for multiple certifications.
    """

    def __init__(
        self,
        document: docx.document.Document,
        certifications: Certifications,
        settings: ResumeCertificationsSettings,
    ):
        """Initialize the basic certifications renderer.

        Args:
            document (docx.document.Document): The Word document object to render into.
            certifications (Certifications): A list of certification data to render.
            settings (ResumeCertificationsSettings): The rendering settings for the certifications section.

        Notes:
            1. Initializes the parent class with the provided document, certifications, and settings.
        """
        super().__init__(document, certifications, settings)

    def render(self) -> None:
        """Render the certifications section of a document.

        Args:
            None

        Returns:
            None: This method modifies the document in-place and does not return a value.

        Notes:
            1. Log the start of the certifications section rendering.
            2. If there are no certifications, return early without adding anything.
            3. Add a level-2 heading titled "Certifications".
            4. For each certification in the list:
               a. Create a new RenderCertificationSection instance.
               b. Call the render method on that instance to add the certification to the document.
            5. This function modifies the document in-place and does not write to disk.
        """
        log.info("Rendering Certifications section.")

        if len(self.certifications) > 0:
            self.document.add_heading("Certifications", level=2)
        else:
            return

        for _certification in self.certifications:
            RenderCertificationSection(
                self.document,
                _certification,
                self.settings,
            ).render()

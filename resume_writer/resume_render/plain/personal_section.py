import logging

import docx.document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from resume_writer.models.personal import ContactInfo, Personal
from resume_writer.resume_render.docx_hyperlink import add_hyperlink
from resume_writer.resume_render.render_settings import ResumePersonalSettings
from resume_writer.resume_render.resume_render_base import ResumeRenderPersonalBase

log = logging.getLogger(__name__)


class RenderPersonalSection(ResumeRenderPersonalBase):
    """Render personal contact info section.

    This class is responsible for rendering the personal information section of a resume,
    including contact details, banner, note, websites, and visa status.

    Attributes:
        document (docx.document.Document): The Word document to which the personal section will be added.
        personal (Personal): The personal information object containing contact details, banner, note, websites, and visa status.
        settings (ResumePersonalSettings): The settings object that controls which elements of the personal section are rendered.
    """

    def __init__(
        self,
        document: docx.document.Document,
        personal: Personal,
        settings: ResumePersonalSettings,
    ):
        """Initialize the personal section renderer.

        Parameters
        ----------
        document : docx.document.Document
            The Word document object to which the personal section will be added.
        personal : Personal
            The personal information object containing contact details, banner, note, websites, and visa status.
        settings : ResumePersonalSettings
            The settings object that controls which elements of the personal section are rendered.

        Notes
        -----
        1. Initialize the base class with the provided document, personal data, and settings.
        2. Log a debug message indicating initialization is complete.

        """
        log.debug("Initializing personal basic render object")
        super().__init__(document, personal, settings)

    def _contact_info(self) -> None:
        """Render the contact information section of the resume.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal and settings attributes.

        Notes
        -----
        1. Log a debug message indicating the section being rendered.
        2. Extract the contact information from the personal attribute.
        3. If a name is present in both the contact info and settings, add it as a heading.
        4. Determine which contact details to render based on the settings.
        5. If any contact details are to be rendered, add a new paragraph to the document.
        6. If the email is to be rendered, add it as a hyperlink to the paragraph.
        7. If the phone number is to be rendered, add it to the paragraph.
        8. If the location is to be rendered, add it to the paragraph.

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        log.debug("Rendering contact info section")

        _info: ContactInfo = self.personal.contact_info
        if _info.name and self.settings.name:
            _heading = self.document.add_heading(_info.name, level=1)
            _heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _render_email = False
        _render_phone = False
        _render_location = False

        _has_content = False

        if _info.email and self.settings.email:
            _render_email = True

        if _info.phone and self.settings.phone:
            _render_phone = True

        if _info.location and self.settings.location:
            _render_location = True

        if _render_email or _render_phone or _render_location:
            _paragraph = self.document.add_paragraph()
            _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if _render_email:
            add_hyperlink(_paragraph, _info.email, f"mailto: {_info.email}")
            _has_content = True

        if _render_phone:
            if _has_content:  # we already have something on this line
                _paragraph.add_run(" | ")
            _paragraph.add_run(_info.phone)
            _has_content = True

        if _render_location:
            if _has_content:  # we already have something on this line
                _paragraph.add_run(" | ")
            _paragraph.add_run(_info.location)

    def _banner(self) -> None:
        """Render the banner section.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal and settings attributes.

        Notes
        -----
        1. Log a debug message indicating the section being rendered.
        2. Retrieve the banner text from the personal object.
        3. If the banner text is not empty, add a heading and the banner text to the document.

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        log.debug("Rendering banner section")

        _banner = self.personal.banner
        if _banner.text:
            self.document.add_heading("Banner", level=3)
            self.document.add_paragraph(_banner.text)

    def _note(self) -> None:
        """Render the note section.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal and settings attributes.

        Notes
        -----
        1. Log a debug message indicating the section being rendered.
        2. Retrieve the note text from the personal object.
        3. If the note text is not empty, add a heading and the note text to the document.

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        log.debug("Rendering note section")

        _note = self.personal.note
        if _note.text:
            self.document.add_heading("Note", level=3)
            self.document.add_paragraph(_note.text)

    def _websites(self) -> str | None:
        """Render the websites section of a resume.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal and settings attributes.

        Notes
        -----
        1. Log a debug message indicating that the websites section is being rendered.
        2. Retrieve the user's websites from the personal object.
        3. Initialize an empty string for the paragraph text and a new paragraph object.
        4. Check if the user's GitHub profile is provided and if the GitHub setting is enabled. If so, add a hyperlink to the GitHub profile to the paragraph and set the has_content flag to True.
        5. Repeat step 4 for the LinkedIn, website, and Twitter profiles.
        6. If the paragraph has content, set its alignment to center.

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        # TODO: This runs over one line if all the sites are populated

        log.debug("Rendering websites section")

        _websites = self.personal.websites

        _paragraph_text = ""
        _paragraph = self.document.add_paragraph()

        _has_content = False

        if _websites.github and self.settings.github:
            add_hyperlink(_paragraph, "GitHub", _websites.github)
            _has_content = True

        if _websites.linkedin and self.settings.linkedin:
            if _has_content:
                _paragraph.add_run(" | ")
            add_hyperlink(_paragraph, "LinkedIn", _websites.linkedin)
            _has_content = True

        if _websites.website and self.settings.website:
            if _has_content:
                _paragraph.add_run(" | ")
            add_hyperlink(_paragraph, "Website", _websites.website)
            _has_content = True

        if _websites.twitter and self.settings.twitter:
            if _has_content:
                _paragraph.add_run(" | ")
            add_hyperlink(_paragraph, "X/Twitter", _websites.twitter)

        if _has_content:
            _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return None

    def _visa_status(self) -> None:
        """Render the visa status section.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal and settings attributes.

        Notes
        -----
        1. Create a new paragraph and set its alignment to center.
        2. Retrieve the visa status object from the personal data.
        3. If work authorization is enabled and present, add it to the paragraph.
        4. If sponsorship requirement is enabled, determine if the value is "Yes" or "No" and add the appropriate text.
        5. If both work authorization and sponsorship are present, add a separator "|".

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        _paragraph = self.document.add_paragraph()
        _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _visa_status = self.personal.visa_status

        if _visa_status.work_authorization and self.settings.work_authorization:
            _work_authorization = _visa_status.work_authorization
            _paragraph.add_run(_work_authorization)

        if _visa_status.require_sponsorship and self.settings.require_sponsorship:
            # handle the case where the value is simply "yes",
            # otherwise use the verbatim value
            _value = "Yes" if _visa_status.require_sponsorship else "No"
            _sponsorship_run = _paragraph.add_run()
            _sponsorship_run.add_text(f"Requires Sponsorship: {_value}")

            # if work authorization was rendered, add a space.
            if _visa_status.work_authorization and self.settings.work_authorization:
                _sponsorship_run.add_text(" | ")

            _sponsorship_run.add_text(_value)

    def render(self) -> None:
        """Render the personal section of a document.

        Parameters
        ----------
        self : RenderPersonalSection
            The instance of the class containing the personal information.

        Notes
        -----
        1. Log a debug message indicating the start of personal section rendering.
        2. Initialize an empty list for banner lines.
        3. Add contact information to the document if enabled in settings.
        4. Add websites to the document if enabled in settings.
        5. Append banner text to banner lines if enabled in settings and not empty.
        6. Append note text to banner lines if enabled in settings and not empty.
        7. Add banner lines to the document if any exist.
        8. Add visa status to the document if enabled in settings.

        Returns
        -------
        None
            The method modifies the document in-place.

        """
        log.debug("Rendering personal section")

        _banner_lines = []

        if self.personal.contact_info and self.settings.contact_info:
            self._contact_info()

        if self.personal.banner and self.settings.banner and self.personal.banner.text:
            _banner_lines.append(self.personal.banner.text)

        if self.personal.note and self.settings.note and self.personal.note.text:
            _banner_lines.append(self.personal.note.text)

        if len(_banner_lines) > 0:
            _banner_paragraph = self.document.add_paragraph()

            _txt = "\n".join(_banner_lines)
            _banner_paragraph.add_run(_txt)
            _banner_paragraph.paragraph_format.space_before = Pt(4)
            _banner_paragraph.paragraph_format.space_after = Pt(10)
            # _banner_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.personal.websites and self.settings.websites:
            self._websites()

        if self.personal.visa_status and self.settings.visa_status:
            self._visa_status()

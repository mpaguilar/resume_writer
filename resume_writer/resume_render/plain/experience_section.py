import logging
import re
from datetime import datetime

import docx
import docx.document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)
from resume_writer.resume_render.docx_hyperlink import add_hyperlink
from resume_writer.resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)
from resume_writer.utils.skills_splitter import skills_splitter

log = logging.getLogger(__name__)

_punctuation_end_re = re.compile(r"[\).,!?;:\]]")


class RenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section.

    Attributes:
        document (docx.document.Document): The DOCX document object to render into.
        role (Role): The role data to render.
        settings (ResumeRolesSettings): The rendering settings for roles.
        font_size (int): The base font size for rendering.
        errors (list[str]): List of error messages encountered during rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document (docx.document.Document): The DOCX document object to render into.
            role (Role): The role data to render.
            settings (ResumeRolesSettings): The rendering settings for roles.

        Returns:
            None

        Notes:
            1. Initializes the base class with the provided document, role, and settings.
            2. Logs a debug message indicating initialization has started.
        """
        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> list[str]:
        """Render role skills section.

        Args:
            None

        Returns:
            A list containing strings of skills. The list is empty if no skills are present.

        Notes:
            1. Creates a new paragraph in the document.
            2. Checks if skills are present in the role and if the settings allow displaying skills.
            3. If skills are present, joins them with commas and adds them as italicized text.
            4. Sets the font size to be two points smaller than the base font size.
            5. Logs a debug message about the skills rendering process.
        """
        log.debug("Rendering role skills.")

        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _paragraph = self.document.add_paragraph()
            _skills_str = ", ".join(self.role.skills)
            _skills_run = _paragraph.add_run(f"Skills: {_skills_str}")
            _skills_run.font.size = Pt(self.font_size - 2)
            _skills_run.font.italic = True

    def _details(self) -> list[str]:
        """Render role details section.

        Args:
            None

        Returns:
            A list of strings containing details like job category, location, agency name, and employment type.

        Notes:
            1. Initializes an empty list to store the details.
            2. Checks if job category is present and if it should be displayed.
            3. Adds job category to the list if applicable.
            4. Checks if location is present and if it should be displayed.
            5. Adds location to the list if applicable.
            6. Checks if agency name is present and if it should be displayed.
            7. Adds agency name to the list if applicable.
            8. Checks if employment type is present and if it should be displayed.
            9. Adds employment type to the list if applicable.
            10. Returns the list of details.
        """
        log.debug("Rendering role details.")
        _paragraph_lines = []

        # job category
        _basics = self.role.basics
        if _basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {_basics.job_category}")

        if _basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {_basics.location}")

        if _basics.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"Agency: {_basics.agency_name}")

        if _basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {_basics.employment_type}")

        return _paragraph_lines

    def _title_and_company(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Render role title and company section.

        Args:
            paragraph (docx.text.paragraph.Paragraph): The paragraph object to add the title and company to.

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of title and company rendering.
            2. Retrieves the basics information from the role.
            3. Checks if the title is present; if not, raises a ValueError with an error message.
            4. Adds the title as bold and underlined text to the paragraph.
            5. Sets the font size to be two points larger than the base font size.
            6. Checks if the company name is present; if not, raises a ValueError with an error message.
            7. Adds the company name with a tab character and bold formatting to the paragraph.
            8. Sets the font size to be two points larger than the base font size.
        """
        log.debug("Rendering role title and company.")

        _basics = self.role.basics

        # title is required
        if not _basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.error(_msg)
            raise ValueError(_msg)

        _title_run = paragraph.add_run()
        _title_run.add_text(f"{_basics.title}")
        _title_run.bold = True
        _title_run.underline = True
        _title_run.font.size = Pt(self.font_size + 2)

        if self.settings.employment_type and _basics.employment_type:
            _employment_type_run = paragraph.add_run(f" ({_basics.employment_type})")
            _employment_type_run.font.size = Pt(self.font_size - 1)

        # company name is required
        if not _basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.error(_msg)
            raise ValueError(_msg)

        _company_run = paragraph.add_run(f"\t{_basics.company}")
        _company_run.bold = True
        _company_run.font.size = Pt(self.font_size + 2)


    def _dates_and_location(
        self,
        paragraph: docx.text.paragraph.Paragraph,
    ) -> list[str]:
        """Render role dates section.

        Args:
            paragraph (docx.text.paragraph.Paragraph): The paragraph object to add the dates and location to.

        Returns:
            A list of strings containing the formatted date and location information.

        Notes:
            1. Logs a debug message indicating the start of date and location rendering.
            2. Retrieves the basics information from the role.
            3. Creates a new run object for the date information.
            4. Checks if the start date is present; if not, logs a warning message.
            5. Formats the start date as "Month Year" and adds it to the run.
            6. If an end date is present, formats it as "Month Year" and adds it to the run with a hyphen.
            7. If no end date is present, adds "Present" to the run.
            8. If location is present and should be displayed, adds it to the paragraph with a tab.
            9. Returns the list of formatted date and location strings.
        """
        log.debug("Rendering role dates.")

        _basics = self.role.basics

        _date_run = paragraph.add_run()

        # Start date
        if not _basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(_basics.start_date, "%B %Y")
        _date_run.add_text(f"{_value}")

        # End date
        if _basics.end_date:
            _value = datetime.strftime(_basics.end_date, "%B %Y")
            _date_run.add_text(f" - {_value}")
        else:
            _date_run.add_text(" - Present")

        if _basics.location and self.settings.location:
            paragraph.add_run(f"\t{_basics.location}")

    def _render_task(
        self,
        paragraph: docx.text.paragraph.Paragraph,
        task_line: str,
    ) -> None:
        """Render a single task line with skill highlighting.

        Args:
            paragraph (docx.text.paragraph.Paragraph): The paragraph object to add the task to.
            task_line (str): The task text to render.

        Returns:
            None

        Notes:
            1. Checks if skills are present and if the settings allow including tasks.
            2. Splits the task line using the skills splitter function.
            3. Initializes leading and trailing space flags.
            4. Iterates through each fragment in the task line.
            5. Creates a new run for each fragment.
            6. Determines if a trailing space is needed based on punctuation.
            7. If the fragment is a skill, adds it in bold.
            8. Adds leading space if needed.
            9. Adds the fragment text to the run.
            10. Adds trailing space if needed.
            11. Adds a line break after the last fragment.
        """
        if (
            self.role.skills
            and self.settings.include_tasks
            and len(self.role.skills) > 0
        ):
            _fragments = skills_splitter(task_line, self.role.skills)
            _leading_space = True
            _trailing_space = True
            for _ndx, _fragment in enumerate(_fragments):
                _run = paragraph.add_run()

                if (_ndx + 1) < len(_fragments) and not _punctuation_end_re.match(
                    _fragments[_ndx + 1],
                ):
                    _trailing_space = True
                else:
                    _trailing_space = False

                if _fragment in self.role.skills:
                    if _leading_space:
                        _run.add_text(" ")

                    _run.add_text(_fragment)
                    _run.bold = True

                    if _trailing_space:
                        _run.add_text(" ")
                else:
                    _run.add_text(_fragment)

                # used on the next iteration of the loop
                _leading_space = not re.search(r"[({[]$", _fragment)

            # _run.add_break()

    def _responsibilities(self) -> None:
        """Render role responsibilities section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of responsibilities rendering.
            2. Creates two paragraphs: one for situation and one for tasks.
            3. Sets spacing before and after both paragraphs.
            4. Splits the responsibilities text by double newlines and then by single newlines.
            5. Iterates through each line in the responsibilities.
            6. If the line starts with "* ", renders it as a task if tasks are included.
            7. If the line doesn't start with "* ", adds it to the situation paragraph if situation is included.
            8. Adds line breaks as needed.
        """
        _msg = f"Rendering responsibilities for role {self.role.basics.title}"
        log.debug(_msg)

        # _situation_paragraph = self.document.add_paragraph()
        # _situation_paragraph.paragraph_format.space_before = Pt(0)
        # _situation_paragraph.paragraph_format.space_after = Pt(6)

        _responsibilities_lines: list[str] = self.role.responsibilities.text.replace(
            "\n\n",
            "\n",
        ).split("\n")

        if len(_responsibilities_lines) > 0:
            _responsibilities_paragraph = self.document.add_paragraph()
            _responsibilities_paragraph.paragraph_format.space_before = Pt(6)
            # _responsibilities_paragraph.paragraph_format.space_after = Pt(6)

            for _line in _responsibilities_lines:
                # task lines start with "* "
                if _line.startswith("* ") and self.settings.include_tasks:
                    self._render_task(_responsibilities_paragraph, _line)
                    if _line is not _responsibilities_lines[-1]:
                        _responsibilities_paragraph.add_run().add_break()

                # if self.settings.include_situation and _line and not _line.startswith("* "):
                #     _situation_run = _situation_paragraph.add_run()
                #     _situation_run.add_text(_line)
                #     _situation_run.add_break()

    def _description(self) -> None:
        """Render role summary and details section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Checks if a summary is present and if the settings allow displaying it.
            2. If a summary is present, creates a new paragraph and adds the summary text in bold and italic.
            3. Sets the spacing after the paragraph.
            4. Sets the spacing before the paragraph.
            5. Checks if responsibilities are present and if the settings allow displaying them.
            6. If responsibilities are present, calls the responsibilities rendering method.
        """
        if self.role.summary and self.settings.summary:
            _summary_paragraph = self.document.add_paragraph()
            _summary_run = _summary_paragraph.add_run()
            _summary_run.add_text(self.role.summary.summary.strip())
            _summary_run.font.italic = True
            _summary_run.font.bold = True
            # _summary_paragraph.paragraph_format.space_after = Pt(self.font_size)
            _summary_paragraph.paragraph_format.space_before = Pt(0)

        if self.role.responsibilities and self.settings.responsibilities:
            self._responsibilities()

    def render(self) -> None:
        """Render role overview/basics section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of rendering.
            2. Creates a paragraph for basics information.
            3. Sets spacing after the paragraph.
            4. Renders the title and company information.
            5. Adds tab stops to the basics paragraph for formatting.
            6. Creates a paragraph for dates and location.
            7. Adds tab stops to the dates paragraph for formatting.
            8. Renders the dates and location information.
            9. Sets spacing after the dates paragraph.
            10. Renders the description section.
            11. Renders the skills section.
        """
        log.debug("Rendering roles section.")

        _basics_paragraph = self.document.add_paragraph()
        _basics_paragraph.paragraph_format.space_after = Pt(self.font_size / 4)
        _basics_paragraph.paragraph_format.space_before = Pt(self.font_size)
        self._title_and_company(_basics_paragraph)

        # add tab stops to format title, company, dates, and location neatly
        _tab_stop_right = Inches(7.4)
        _tab_stops = _basics_paragraph.paragraph_format.tab_stops

        _tab_stops.add_tab_stop(
            _tab_stop_right,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )

        _dates_paragraph = self.document.add_paragraph()
        # add tab stops to format title, company, dates, and location neatly
        _tab_stop_right = Inches(7.4)
        _tab_stops = _dates_paragraph.paragraph_format.tab_stops

        _tab_stops.add_tab_stop(
            _tab_stop_right,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )

        self._dates_and_location(_dates_paragraph)
        _dates_paragraph.paragraph_format.space_after = Pt(self.font_size / 2)

        self._description()
        self._skills()


class RenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section.

    Attributes:
        document (docx.document.Document): The DOCX document object to render into.
        roles (Roles): The list of role data to render.
        settings (ResumeRolesSettings): The rendering settings for roles.
        font_size (int): The base font size for rendering.
        errors (list[str]): List of error messages encountered during rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document (docx.document.Document): The DOCX document object to render into.
            roles (Roles): The list of role data to render.
            settings (ResumeRolesSettings): The rendering settings for roles.

        Returns:
            None

        Notes:
            1. Initializes the base class with the provided document, roles, and settings.
            2. Logs a debug message indicating initialization has started.
        """
        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of rendering.
            2. Checks if there are any roles to render.
            3. If roles exist, adds a centered heading with the text "Work History".
            4. If no roles exist, logs an info message.
            5. Iterates through each role and renders it using the RenderRoleSection class.
            6. Adds a blank paragraph after each role with 12 points of spacing after.
            7. Adds a horizontal line after each role.
        """
        log.debug("Rendering roles section.")
        if len(self.roles) > 0:
            _heading = self.document.add_heading("Work History", level=2)
            _heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            log.info("No roles found")
            return

        for _role in self.roles:
            RenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            ).render()



class RenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section.

    Attributes:
        document (docx.document.Document): The DOCX document object to render into.
        project (Project): The project data to render.
        settings (ResumeProjectsSettings): The rendering settings for projects.
        font_size (int): The base font size for rendering.
        errors (list[str]): List of error messages encountered during rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object.

        Args:
            document (docx.document.Document): The DOCX document object to render into.
            project (Project): The project data to render.
            settings (ResumeProjectsSettings): The rendering settings for projects.

        Returns:
            None

        Notes:
            1. Initializes the base class with the provided document, project, and settings.
            2. Logs a debug message indicating initialization has started.
        """
        log.debug("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> None:
        """Render project overview section.

        Args:
            paragraph (docx.text.paragraph.Paragraph): The paragraph object to add the overview to.

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of overview rendering.
            2. Retrieves the overview information from the project.
            3. Adds the project title as bold text to the paragraph.
            4. Adds a tab character to the paragraph.
            5. Sets spacing after the paragraph.
            6. Adds a hyperlink to the project website using the add_hyperlink function.
        """
        log.debug("Rendering project overview.")

        _overview_paragraph = self.document.add_paragraph()
        # add tab stops to format title, company, dates, and location neatly
        _tab_stop_right = Inches(7.4)
        _tab_stops = _overview_paragraph.paragraph_format.tab_stops

        _tab_stops.add_tab_stop(
            _tab_stop_right,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )

        _overview = self.project.overview

        _title_run = _overview_paragraph.add_run()
        _title_run.add_text(_overview.title)
        _title_run.bold = True
        _title_run.underline = True
        _title_run.font.size = Pt(self.font_size + 2)

        _overview_paragraph.add_run("\t")
        _overview_paragraph.paragraph_format.space_after = Pt(self.font_size / 2)
        _overview_paragraph.paragraph_format.space_before = Pt(self.font_size / 2)

        self.add_hyperlink(_overview_paragraph, f"{_overview.url_description}", _overview.url)

    def _skills(self, paragraph: docx.text.paragraph.Paragraph) -> list[str]:
        """Render project skills section.

        Args:
            paragraph (docx.text.paragraph.Paragraph): The paragraph object to add the skills to.

        Returns:
            A list of strings containing the project skills.

        Notes:
            1. Logs a debug message indicating the start of skills rendering.
            2. Joins the project skills with commas.
            3. Creates a new run object for the skills text.
            4. Adds the skills text with "Skills: " prefix.
            5. Sets the text to be italic.
            6. Returns the list of skills.
        """
        log.debug("Rendering project skills.")

        _skills = ", ".join(self.project.skills)
        _skills_run = paragraph.add_run()
        _skills_run.font.size = Pt(self.font_size - 2)
        _skills_run.add_text(f"Skills: {_skills}")
        _skills_run.italic = True

    def render(self) -> None:
        """Render project section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of project rendering.
            2. Creates a paragraph for the project information.
            3. Adds tab stops to the paragraph for formatting.
            4. Checks if the overview should be displayed and if it exists.
            5. If the overview should be displayed, renders the overview section.
            6. Sets spacing after the overview paragraph.
            7. Checks if the description should be displayed and if it exists.
            8. If the description should be displayed, adds each line of the description as a separate run with line breaks.
            9. Sets spacing before and after the description paragraph.
            10. Checks if skills should be displayed and if they exist.
            11. If skills should be displayed, adds a new paragraph for skills and renders them.
        """
        log.debug("Rendering project section.")

        if self.settings.overview and self.project.overview:
            self._overview()

        if self.settings.description and self.project.description:
            _description_paragraph = self.document.add_paragraph()
            self._highlight_skills(_description_paragraph)

            _description_paragraph.paragraph_format.space_before = Pt(0)
            _description_paragraph.paragraph_format.space_after = Pt(0)

        if self.settings.skills and self.project.skills:
            _skills_paragraph = self.document.add_paragraph()
            self._skills(_skills_paragraph)
            _skills_paragraph.paragraph_format.space_before = Pt(0)
            _skills_paragraph.paragraph_format.space_after = Pt(0)

    def _highlight_skills(self, paragraph : docx.text.paragraph.Paragraph) -> None:
        if(
            self.project.skills
            and len(self.project.skills) > 0
        ):
            _fragments = skills_splitter(self.project.description.text, self.project.skills)
            _leading_space = True
            _trailing_space = True
            for _ndx, _fragment in enumerate(_fragments):
                _run = paragraph.add_run()

                if (_ndx + 1) < len(_fragments) and not _punctuation_end_re.match(
                    _fragments[_ndx + 1],
                ):
                    _trailing_space = True
                else:
                    _trailing_space = False

                if _fragment in self.project.skills:
                    if _leading_space:
                        _run.add_text(" ")

                    _run.add_text(_fragment)
                    _run.bold = True

                    if _trailing_space:
                        _run.add_text(" ")
                else:
                    _run.add_text(_fragment)

                # used on the next iteration of the loop
                _leading_space = not re.search(r"[({[]$", _fragment)


class RenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section.

    Attributes:
        document (docx.document.Document): The DOCX document object to render into.
        projects (Projects): The list of project data to render.
        settings (ResumeProjectsSettings): The rendering settings for projects.
        font_size (int): The base font size for rendering.
        errors (list[str]): List of error messages encountered during rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object.

        Args:
            document (docx.document.Document): The DOCX document object to render into.
            projects (Projects): The list of project data to render.
            settings (ResumeProjectsSettings): The rendering settings for projects.

        Returns:
            None

        Notes:
            1. Initializes the base class with the provided document, projects, and settings.
            2. Logs a debug message indicating initialization has started.
        """
        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of projects rendering.
            2. Checks if there are any projects to render.
            3. If projects exist, adds a heading with the text "Projects".
            4. Iterates through each project and renders it using the RenderProjectSection class.
        """
        log.debug("Rendering projects section.")
        if len(self.projects) > 0:
            _heading = self.document.add_heading("Projects", level=2)
            _heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            log.info("No projects found")
            return

        for _project in self.projects:
            RenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()


class RenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section.

    Attributes:
        document (docx.document.Document): The DOCX document object to render into.
        experience (Experience): The experience data to render.
        settings (ResumeExperienceSettings): The rendering settings for experience.
        font_size (int): The base font size for rendering.
        errors (list[str]): List of error messages encountered during rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document (docx.document.Document): The DOCX document object to render into.
            experience (Experience): The experience data to render.
            settings (ResumeExperienceSettings): The rendering settings for experience.

        Returns:
            None

        Notes:
            1. Initializes the base class with the provided document, experience, and settings.
            2. Logs a debug message indicating initialization has started.
        """
        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Define a function to render the roles section, which is called if settings permit.
            2. Define a function to render the projects section, which is called if settings permit.
            3. If `render_projects_first` is `True` in settings, call project render function then role render function.
            4. Otherwise, call role render function then project render function.
        """
        log.debug("Rendering experience section.")

        def _render_roles() -> None:
            if self.settings.roles and self.experience.roles:
                RenderRolesSection(
                    document=self.document,
                    roles=self.experience.roles,
                    settings=self.settings.roles_settings,
                ).render()

        def _render_projects() -> None:
            if (
                self.settings.projects
                and self.experience.projects
                and len(self.experience.projects) > 0
            ):
                RenderProjectsSection(
                    document=self.document,
                    projects=self.experience.projects,
                    settings=self.settings.projects_settings,
                ).render()

        if self.settings.render_projects_first:
            _render_projects()
            # self.document.add_paragraph().paragraph_format.space_after = Pt(4)
            _render_roles()
        else:
            _render_roles()
            # self.document.add_paragraph().paragraph_format.space_after = Pt(4)
            _render_projects()

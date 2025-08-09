import logging

import docx.document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt

from resume_writer.models.education import Degree, Education
from resume_writer.resume_render.render_settings import ResumeEducationSettings
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderDegreeBase,
    ResumeRenderEducationBase,
)

log = logging.getLogger(__name__)


class RenderDegreeSection(ResumeRenderDegreeBase):
    """Render a single academic degree in a resume document.

    This class handles the visual formatting and layout of a degree entry,
    including school name, degree name, dates, and GPA, using a structured
    paragraph with tab stops for alignment.

    Inherits from:
        ResumeRenderDegreeBase: Base class providing shared rendering logic.

    Attributes:
        document (docx.document.Document): The Word document being rendered.
        degree (Degree): The academic degree data to render.
        settings (ResumeEducationSettings): Configuration settings for rendering.
        font_size (int): Base font size used for text (in points).
    """

    def __init__(
        self,
        document: docx.document.Document,
        degree: Degree,
        settings: ResumeEducationSettings,
    ):
        """Initialize the degree renderer.

        Args:
            document (docx.document.Document): The DOCX document to render into.
            degree (Degree): The academic degree data to render.
            settings (ResumeEducationSettings): Rendering configuration for degrees.

        Returns:
            None

        Notes:
            1. Calls the parent constructor with the provided document, degree, and settings.
        """
        super().__init__(document=document, degree=degree, settings=settings)

    def render(self) -> None:
        """Render a single academic degree entry into the document.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of degree rendering.
            2. Creates a new paragraph in the document.
            3. Adds a right-aligned tab stop at 7.4 inches to align text elements.
            4. If the school name is provided and enabled in settings, adds it in bold with a larger font size.
            5. If the degree name is provided and enabled in settings, adds it after a tab, in bold, followed by a line break.
            6. If the start date is provided and enabled in settings, formats it as "Month Year" and adds it to the paragraph.
            7. If the end date is provided and enabled in settings, formats it as "Month Year" and adds it with a dash to the paragraph.
            8. If the GPA is provided and enabled in settings, adds it after a tab with the label "GPA: ".
        """
        log.debug("Rendering degree section.")

        _paragraph = self.document.add_paragraph()

        # add tab stops to format title, company, dates, and location neatly
        _tab_stop_right = Inches(7.4)
        _tab_stops = _paragraph.paragraph_format.tab_stops

        _tab_stops.add_tab_stop(
            _tab_stop_right,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.SPACES,
        )

        # School
        if self.degree.school and self.settings.school:
            _school_run = _paragraph.add_run(f"{self.degree.school}")
            _school_run.bold = True
            _school_run.font.size = Pt(self.font_size + 2)

        if self.degree.degree and self.settings.degree:
            _degree_run = _paragraph.add_run()
            _degree_run.add_text(f"\t{self.degree.degree}")
            _degree_run.bold = True
            _degree_run.add_break()

        if self.degree.start_date and self.settings.start_date:
            _value = self.degree.start_date.strftime("%B %Y")
            _start_date_run = _paragraph.add_run()
            _start_date_run.add_text(f"{_value}")

        if self.degree.end_date and self.settings.end_date:
            _value = self.degree.end_date.strftime("%B %Y")
            _end_date_run = _paragraph.add_run()
            _end_date_run.add_text(f" - {_value}")

        if self.degree.gpa and self.settings.gpa:
            _degree_run = _paragraph.add_run(f"\tGPA: {self.degree.gpa}")


class RenderEducationSection(ResumeRenderEducationBase):
    """Render the entire education section of a resume.

    This class manages the rendering of the education section, including
    a centered heading and individual degree entries formatted via
    RenderDegreeSection.

    Inherits from:
        ResumeRenderEducationBase: Base class providing shared rendering logic.

    Attributes:
        document (docx.document.Document): The Word document being rendered.
        education (Education): The education data to render.
        settings (ResumeEducationSettings): Configuration settings for rendering.
    """

    def __init__(
        self,
        document: docx.document.Document,
        education: Education,
        settings: ResumeEducationSettings,
    ):
        """Initialize the education renderer.

        Args:
            document (docx.document.Document): The DOCX document to render into.
            education (Education): The education data to render.
            settings (ResumeEducationSettings): Rendering configuration for education.

        Returns:
            None

        Notes:
            1. Calls the parent constructor with the provided document, education, and settings.
        """
        super().__init__(document, education, settings)

    def render(self) -> None:
        """Render the education section into the document.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of education rendering.
            2. Exits early if degree rendering is disabled in settings.
            3. Adds a centered heading titled "Education" with level 2.
            4. Iterates through each degree in the education data.
            5. For each degree, creates a RenderDegreeSection instance and calls its render method.
            6. Adds a blank paragraph between degrees to provide visual separation.
        """
        log.debug("Rendering education section.")

        if not self.settings.degrees:
            return

        _heading = self.document.add_heading("Education", level=2)
        _heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _degree in self.education.degrees:
            RenderDegreeSection(
                document=self.document,
                degree=_degree,
                settings=self.settings,
            ).render()
            # add a blank line between degrees
            self.document.add_paragraph()

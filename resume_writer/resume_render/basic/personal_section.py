import logging

import docx.document
from resume_render.render_settings import ResumePersonalSettings
from resume_render.resume_render_base import ResumeRenderPersonalBase

from resume_writer.models.personal import ContactInfo, Personal

log = logging.getLogger(__name__)


class RenderPersonalSection(ResumeRenderPersonalBase):
    """Render personal contact info section."""

    def __init__(
        self,
        document: docx.document.Document,
        personal: Personal,
        settings: ResumePersonalSettings,
    ):
        """Initialize the personal section renderer.

        Args:
            document: The Word document object to which content will be added.
            personal: The personal information model containing contact details, banner, note, websites, and visa status.
            settings: The rendering settings that control which sections are enabled or disabled.

        Returns:
            None

        Notes:
            1. Initialize the base class with the provided document, personal data, and settings.
            2. Log debug message indicating the initialization of the personal basic render object.

        """
        log.debug("Initializing personal basic render object")
        super().__init__(document, personal, settings)

    def _contact_info(self) -> None:
        """Render the contact info section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Initialize an empty list to hold lines of contact information.
            2. Extract the contact info from the personal data.
            3. If the name is present and enabled in settings, add "Name: <name>" to the list.
            4. If the email is present and enabled in settings, add "Email: <email>" to the list.
            5. If the phone is present and enabled in settings, add "Phone: <phone>" to the list.
            6. If the location is present and enabled in settings, add "Location: <location>" to the list.
            7. If any lines were added, join them with newlines and add the combined text as a paragraph in the document.

        """
        log.debug("Rendering contact info section")

        _paragraph_lines = []

        _info: ContactInfo = self.personal.contact_info
        if _info.name and self.settings.name:
            _paragraph_lines.append(f"Name: {_info.name}")

        if _info.email and self.settings.email:
            _paragraph_lines.append(f"Email: {_info.email}")

        if _info.phone and self.settings.phone:
            _paragraph_lines.append(f"Phone: {_info.phone}")

        if _info.location and self.settings.location:
            _paragraph_lines.append(f"Location: {_info.location}")

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def _banner(self) -> None:
        """Render the banner section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Extract the banner text from the personal data.
            2. If the banner text is present, add a level 3 heading titled "Banner".
            3. Add the banner text as a paragraph immediately after the heading.

        """
        log.debug("Rendering banner section")

        _banner = self.personal.banner
        if _banner.text:
            self.document.add_heading("Banner", level=3)
            self.document.add_paragraph(_banner.text)

    def _note(self) -> None:
        """Render the note section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Extract the note text from the personal data.
            2. If the note text is present, add a level 3 heading titled "Note".
            3. Add the note text as a paragraph immediately after the heading.

        """
        log.debug("Rendering note section")

        _note = self.personal.note
        if _note.text:
            self.document.add_heading("Note", level=3)
            self.document.add_paragraph(_note.text)

    def _websites(self) -> None:
        """Render the websites section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Initialize an empty list to hold lines of website information.
            2. Extract the websites data from the personal data.
            3. If GitHub URL is present and enabled in settings, add "GitHub: <url>" to the list.
            4. If LinkedIn URL is present and enabled in settings, add "LinkedIn: <url>" to the list.
            5. If personal website URL is present and enabled in settings, add "Website: <url>" to the list.
            6. If Twitter URL is present and enabled in settings, add "Twitter: <url>" to the list.
            7. If any lines were added, add a level 3 heading titled "Websites".
            8. Join the lines with newlines and add the combined text as a paragraph in the document.

        """
        log.debug("Rendering websites section")

        _paragraph_lines = []

        _websites = self.personal.websites

        if _websites.github and self.settings.github:
            _paragraph_lines.append(f"GitHub: {_websites.github}")

        if _websites.linkedin and self.settings.linkedin:
            _paragraph_lines.append(f"LinkedIn: {_websites.linkedin}")

        if _websites.website and self.settings.website:
            _paragraph_lines.append(f"Website: {_websites.website}")

        if _websites.twitter and self.settings.twitter:
            _paragraph_lines.append(f"Twitter: {_websites.twitter}")

        if len(_paragraph_lines) > 0:
            self.document.add_heading("Websites", level=3)
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def _visa_status(self) -> None:
        """Render the visa status section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Initialize an empty list to hold lines of visa status information.
            2. Extract the visa status data from the personal data.
            3. If work authorization is present and enabled in settings, add "Work Authorization: <status>" to the list.
            4. If sponsorship requirement is defined and enabled in settings, determine the value ("Yes" if True, "No" if False) and add "Require Sponsorship: <value>" to the list.
            5. If any lines were added, add a level 3 heading titled "Visa Status".
            6. Join the lines with newlines and add the combined text as a paragraph in the document.

        """
        _paragraph_lines = []

        _visa_status = self.personal.visa_status

        if _visa_status.work_authorization and self.settings.work_authorization:
            _paragraph_lines.append(
                f"Work Authorization: {_visa_status.work_authorization}",
            )

        if (
            _visa_status.require_sponsorship is not None
            and self.settings.require_sponsorship
        ):
            _value = "Yes" if _visa_status.require_sponsorship else "No"
            _paragraph_lines.append(f"Require Sponsorship: {_value}")

        if len(_paragraph_lines) > 0:
            self.document.add_heading("Visa Status", level=3)
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def render(self) -> None:
        """Render the personal section.

        Args:
            None

        Returns:
            None

        Notes:
            1. If contact info exists and is enabled in settings, render the contact info section.
            2. If banner exists and is enabled in settings, render the banner section.
            3. If note exists and is enabled in settings, render the note section.
            4. If websites exist and are enabled in settings, render the websites section.
            5. If visa status exists and is enabled in settings, render the visa status section.

        """
        if self.personal.contact_info and self.settings.contact_info:
            self._contact_info()

        if self.personal.banner and self.settings.banner:
            self._banner()

        if self.personal.note and self.settings.note:
            self._note()

        if self.personal.websites and self.settings.websites:
            self._websites()

        if self.personal.visa_status and self.settings.visa_status:
            self._visa_status()

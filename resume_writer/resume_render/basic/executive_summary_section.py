import logging
from datetime import datetime

import docx.document

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.resume_render.render_settings import (
    ResumeExecutiveSummarySettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExecutiveSummaryBase,
)
from resume_writer.utils.executive_summary import ExecutiveSummary

log = logging.getLogger(__name__)


class RenderExecutiveSummarySection(ResumeRenderExecutiveSummaryBase):
    """Render experience for a functional resume.

    Inherits from:
        ResumeRenderExecutiveSummaryBase: Base class for rendering executive summaries.

    Attributes:
        document (docx.document.Document): The Docx document object to which the executive summary will be added.
        experience (Experience): The Experience object containing role and job data to summarize.
        settings (ResumeExecutiveSummarySettings): The settings object that defines which categories to include in the summary.
    """

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExecutiveSummarySettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document (docx.document.Document): The Docx document object to which the executive summary will be added.
            experience (Experience): The Experience object containing role and job data to summarize.
            settings (ResumeExecutiveSummarySettings): The settings object that defines which categories to include in the summary.

        Returns:
            None

        Notes:
            1. Initialize the parent class (ResumeRenderExecutiveSummaryBase) with the provided document, experience, and settings.
            2. Log a debug message indicating that the functional experience render object is being initialized.
        """
        log.debug("Initializing functional experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Log a debug message indicating that the functional experience section is being rendered.
            2. Check if the experience object has any roles; if not, raise a ValueError.
            3. Create an ExecutiveSummary object using the experience data.
            4. Generate the executive summary using the specified categories from the settings.
            5. For each category in the summary:
                a. Add a heading to the document with level 4.
                b. For each summary entry in the category:
                    i. Create a new paragraph and apply the "List Bullet" style.
                    ii. Add the summary text as a run.
                    iii. If no company is available, log a warning.
                    iv. If a company is available, determine the date string (either the last date formatted to year or "Present").
                    v. Add an italicized run with the company and date information.
            6. This function performs no disk, network, or database access.
        """
        log.debug("Rendering functional experience section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        _o_executive_summary = ExecutiveSummary(self.experience)
        _executive_summary = _o_executive_summary.summary(self.settings.categories)

        for _category in _executive_summary:
            self.document.add_heading(_category, level=4)
            for _summary in _executive_summary[_category]:
                _paragraph = self.document.add_paragraph()
                _paragraph.style = "List Bullet"

                _paragraph.add_run(_summary["summary"])

                if not _summary["company"]:
                    log.warning(f"No company available for {_summary['title']}")
                else:
                    if _summary["last_date"]:
                        _date_str = datetime.strftime(_summary["last_date"], "%Y")
                    else:
                        _date_str = "Present"

                    _run = _paragraph.add_run(f" ({_summary['company']}, {_date_str})")
                    _run.italic = True

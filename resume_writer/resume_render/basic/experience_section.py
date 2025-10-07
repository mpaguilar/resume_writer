import logging
from datetime import datetime

import docx.document

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)
from resume_writer.resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)

log = logging.getLogger(__name__)


class RenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section.

    This class is responsible for rendering individual job roles in a resume document.
    It formats and adds role details such as company, title, dates, skills, and responsibilities.

    Attributes:
        document (docx.document.Document): The Word document object to which content will be added.
        role (Role): The role object containing job-specific details.
        settings (ResumeRolesSettings): Configuration settings for which role details to render.
        errors (list[str]): List of error messages encountered during rendering.
    """
    
    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document (docx.document.Document): The Word document object to which the role content will be added.
            role (Role): The Role object containing role-specific details such as company, title, dates, etc.
            settings (ResumeRolesSettings): The settings object that controls which parts of the role to render.

        Returns:
            None.

        Notes:
            1. Initialize the base class with the provided document, role, and settings.
            2. Log a debug message indicating initialization.

        """
        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> list[str]:
        """Render role skills section.

        Args:
            None.

        Returns:
            A list of strings, each representing a formatted line for skills. If no skills are present or settings are disabled, returns an empty list.

        Notes:
            1. Initialize an empty list to hold the skill lines.
            2. Check if the role has skills and if the skills setting is enabled.
            3. If both conditions are true, join the skills into a comma-separated string and append it to the list with the label "Skills: ".
            4. Return the list of skill lines.

        """
        log.debug("Rendering role skills.")
        _paragraph_lines = []
        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _skills_str = ", ".join(self.role.skills)
            _paragraph_lines.append(f"Skills: {_skills_str}")

        return _paragraph_lines

    def _details(self) -> list[str]:
        """Render role details section.

        Args:
            None.

        Returns:
            A list of strings, each representing a formatted line for role details such as job category, location, agency, or employment type. Returns an empty list if no details are to be rendered.

        Notes:
            1. Initialize an empty list to store detail lines.
            2. Extract the role's basics information.
            3. Check if job category is present and if the job_category setting is enabled, then append the formatted line.
            4. Repeat for location, agency name, and employment type if applicable and settings are enabled.
            5. Return the list of detail lines.

        """
        log.debug("Rendering role details.")
        _paragraph_lines = []
        # job category
        _basics = self.role.basics
        if _basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {_basics.job_category}")

        if _basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {_basics.location}")

        if _basics.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"Agency: {_basics.agency_name}")

        if _basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {_basics.employment_type}")

        return _paragraph_lines

    def _dates(self) -> list[str]:
        """Render role dates section.

        Args:
            None.

        Returns:
            A list of strings, with a single entry containing the formatted date range (e.g., "01-2020 - 12-2022" or "01-2020 - Present").

        Notes:
            1. Initialize an empty list to hold date lines.
            2. Extract the role's basics information.
            3. Validate that a start date is present; if not, append a warning to errors and return.
            4. Format the start date as "MM-YYYY".
            5. If an end date is present, format it as "MM-YYYY" and append it with a hyphen.
            6. If no end date is present, append " - Present".
            7. Return the formatted date string as a list with one element.

        """
        log.debug("Rendering role dates.")
        _paragraph_lines = []

        _basics = self.role.basics
        # Start date
        if not _basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(_basics.start_date, "%m-%Y")
        _start_and_end = f"{_value}"

        # End date
        if _basics.end_date:
            _value = datetime.strftime(_basics.end_date, "%m-%Y")
            _start_and_end += f" - {_value}"
        else:
            _start_and_end += " - Present"

        return [_start_and_end]

    def render(self) -> None:
        """Render role overview/basics section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Initialize an empty list to hold paragraph lines.
            2. Extract the role's basics information.
            3. Check if the company name is present; if not, append a warning and log it.
            4. Append the formatted "Company: <name>" line.
            5. Render the dates using _dates method and extend the paragraph lines.
            6. Check if the title is present; if not, append a warning and log it.
            7. Append the formatted "Title: <title>" line.
            8. Render the details using _details method and extend the paragraph lines.
            9. Clean the paragraph lines by replacing double newlines with single newlines.
            10. If any paragraph lines exist, add a paragraph to the document with the joined lines.
            11. If a summary exists and summary setting is enabled, add it as a paragraph.
            12. If responsibilities exist and responsibilities setting is enabled, add them as a paragraph with cleaned text.
            13. Render the skills using _skills method and add them as a paragraph if any exist.

        """
        _paragraph_lines = []

        log.debug("Rendering roles section.")
        _basics = self.role.basics
        # company name is required
        if not _basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Company: {_basics.company}")

        _date_lines = self._dates()
        if len(_date_lines) > 0:
            _paragraph_lines.extend(_date_lines)

        if not _basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Title: {_basics.title}")

        _detail_lines = self._details()
        if len(_detail_lines) > 0:
            _paragraph_lines.extend(_detail_lines)

        _clean_lines = [_line.replace("\n\n", "\n") for _line in _paragraph_lines]

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_clean_lines))

        if self.role.summary and self.settings.summary:
            self.document.add_paragraph(self.role.summary.summary)

        if self.role.responsibilities and self.settings.responsibilities:
            _responsibilities_text = self.role.responsibilities.text.replace(
                "\n\n",
                "\n",
            )
            self.document.add_paragraph(_responsibilities_text)

        _skills_lines = self._skills()
        if len(_skills_lines) > 0:
            _skills_paragraph = self.document.add_paragraph()
            _skills_paragraph.add_run("\n".join(_skills_lines))


class RenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section.

    This class is responsible for rendering multiple job roles in a resume document.
    It manages the formatting and ordering of job roles, adding appropriate headings and spacing.

    Attributes:
        document (docx.document.Document): The Word document object to which the roles content will be added.
        roles (Roles): A list of Role objects representing job roles to be rendered.
        settings (ResumeRolesSettings): Configuration settings for which role details to render.
        errors (list[str]): List of error messages encountered during rendering.
    """
    
    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document (docx.document.Document): The Docx document object to which the roles content will be added.
            roles (Roles): A list of Role objects representing job roles to be rendered.
            settings (ResumeRolesSettings): The settings object that controls which parts of the roles to render.

        Returns:
            None.

        Notes:
            1. Initialize the base class with the provided document, roles, and settings.

        """
        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Log a debug message indicating the start of rendering.
            2. Store the roles list to avoid filtering twice.
            3. Add a heading "Work History" at level 2 if any roles exist.
            4. If no roles exist, log an info message and return.
            5. Iterate over each role in the roles list.
            6. For each role, create a RenderRoleSection instance and call its render method.
            7. After each role (except the last), add two blank paragraphs to separate entries.

        """
        log.debug("Rendering roles section.")
        # to prevent filtering twice
        _roles = self.roles

        # don't emit anything if there are no roles
        if len(_roles) > 0:
            self.document.add_heading("Work History", level=2)
        else:
            log.info("No roles found")
            return

        for _role in _roles:
            RenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            ).render()

            # add two blank lines between roles
            if _role != _roles[-1]:
                self.document.add_paragraph()
                self.document.add_paragraph()


class RenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section.

    This class is responsible for rendering individual projects in a resume document.
    It formats and adds project details such as description, skills, URLs, and dates.

    Attributes:
        document (docx.document.Document): The Word document object to which the project content will be added.
        project (Project): The Project object containing project details such as description, skills, URLs, etc.
        settings (ResumeProjectsSettings): Configuration settings for which project details to render.
        errors (list[str]): List of error messages encountered during rendering.
    """
    
    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object.

        Args:
            document (docx.document.Document): The Docx document object to which the project content will be added.
            project (Project): The Project object containing project details such as description, skills, URLs, etc.
            settings (ResumeProjectsSettings): The settings object that controls which parts of the project to render.

        Returns:
            None.

        Notes:
            1. Initialize the base class with the provided document, project, and settings.
            2. Log a debug message indicating initialization.

        """
        log.debug("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> list[str]:
        """Render project overview section.

        Args:
            None.

        Returns:
            A list of strings, each representing a formatted line for the project overview such as URL description, URL, start date, or end date. Returns an empty list if no overview data is to be rendered.

        Notes:
            1. Initialize an empty list to store overview lines.
            2. Extract the project's overview information.
            3. If URL description is enabled and present, format it as a string.
            4. If URL is enabled and present, add it in parentheses to the URL description line.
            5. Strip and add the resulting URL line if it is not empty.
            6. If start date is enabled and present, format the date as "Month Year" and add it.
            7. If end date is enabled and present, format the date as "Month Year" and add it.
            8. Return the list of overview lines.

        """
        log.debug("Rendering project overview.")

        _paragraph_lines = []
        _overview = self.project.overview

        _url_line = ""
        if self.settings.url_description and _overview.url_description:
            _url_line = f"{_overview.url_description} "
        if self.settings.url and _overview.url:
            _url_line += f" ({_overview.url})"

        _url_line = _url_line.strip()
        if _url_line:
            _paragraph_lines.append(_url_line)

        if self.settings.start_date and _overview.start_date:
            _start_date = datetime.strftime(_overview.start_date, "%B %Y")
            _paragraph_lines.append(f"Start Date: {_start_date}")

        if self.settings.end_date and _overview.end_date:
            _end_date = datetime.strftime(_overview.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_end_date}")

        return _paragraph_lines

    def _skills(self) -> list[str]:
        """Render project skills section.

        Args:
            None.

        Returns:
            A list of strings, with a single entry containing "Skills: <comma-separated skills>" if skills are present and enabled. Returns an empty list otherwise.

        Notes:
            1. Initialize an empty list to store skill lines.
            2. Join the project's skills into a comma-separated string.
            3. Append a formatted "Skills: ..." line to the list.
            4. Return the list of skill lines.

        """
        log.debug("Rendering project skills.")

        _paragraph_lines = []

        _skills = ", ".join(self.project.skills)
        _paragraph_lines.append(f"Skills: {_skills}")

        return _paragraph_lines

    def render(self) -> None:
        """Render project section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Log a debug message indicating the start of rendering.
            2. Initialize an empty list to store paragraph lines.
            3. If overview is enabled and the project has an overview, render it using _overview method and extend the paragraph lines.
            4. If description is enabled and the project has a description, append the description text.
            5. If skills are enabled and the project has skills, render them using _skills method and extend the paragraph lines.
            6. If any paragraph lines exist, add a paragraph to the document with the joined lines.

        """
        log.debug("Rendering project section.")

        _paragraph_lines = []

        if self.settings.overview and self.project.overview:
            _paragraph_lines.extend(self._overview())

        if self.settings.description and self.project.description:
            _paragraph_lines.append(self.project.description.text)

        if self.settings.skills and len(self.project.skills) > 0:
            _paragraph_lines.extend(self._skills())

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class RenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section.

    This class is responsible for rendering multiple projects in a resume document.
    It manages the formatting and ordering of projects, adding appropriate headings and spacing.

    Attributes:
        document (docx.document.Document): The Word document object to which the projects content will be added.
        projects (Projects): A list of Project objects representing projects to be rendered.
        settings (ResumeProjectsSettings): Configuration settings for which project details to render.
        errors (list[str]): List of error messages encountered during rendering.
    """
    
    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object.

        Args:
            document (docx.document.Document): The Docx document object to which the projects content will be added.
            projects (Projects): A list of Project objects representing projects to be rendered.
            settings (ResumeProjectsSettings): The settings object that controls which parts of the projects to render.

        Returns:
            None.

        Notes:
            1. Initialize the base class with the provided document, projects, and settings.
            2. Log a debug message indicating initialization.

        """
        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Log a debug message indicating the start of rendering.
            2. If any projects exist, add a heading "Projects" at level 2.
            3. Iterate over each project in the projects list.
            4. For each project, create a RenderProjectSection instance and call its render method.

        """
        log.debug("Rendering projects section.")
        if len(self.projects) > 0:
            self.document.add_heading("Projects", level=2)
        for _project in self.projects:
            _project_render = RenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()


class RenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section.

    This class is responsible for rendering the entire experience section of a resume.
    It coordinates the rendering of roles and projects, applying appropriate formatting and settings.

    Attributes:
        document (docx.document.Document): The Word document object to which the experience content will be added.
        experience (Experience): The Experience object containing roles and projects to be rendered.
        settings (ResumeExperienceSettings): Configuration settings for which experience details to render.
        errors (list[str]): List of error messages encountered during rendering.
    """
    
    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document (docx.document.Document): The Docx document object to which the experience content will be added.
            experience (Experience): The Experience object containing roles and projects to be rendered.
            settings (ResumeExperienceSettings): The settings object that controls which parts of experience to render.

        Returns:
            None.

        Notes:
            1. Initialize the base class with the provided document, experience, and settings.
            2. Log a debug message indicating initialization.

        """
        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section.

        Args:
            None

        Returns:
            None

        Notes:
            1. Define a function to render the roles section, which is called if settings permit.
            2. Define a function to render the projects section, which is called if settings permit.
            3. If `render_projects_first` is `True` in settings, call project render function then role render function.
            4. Otherwise, call role render function then project render function.
        """
        log.debug("Rendering experience section.")

        def _render_roles() -> None:
            if self.settings.roles and self.experience.roles:
                RenderRolesSection(
                    document=self.document,
                    roles=self.experience.roles,
                    settings=self.settings.roles_settings,
                ).render()

        def _render_projects() -> None:
            if self.settings.projects and self.experience.projects:
                RenderProjectsSection(
                    document=self.document,
                    projects=self.experience.projects,
                    settings=self.settings.projects_settings,
                ).render()

        if self.settings.render_projects_first:
            _render_projects()
            _render_roles()
        else:
            _render_roles()
            _render_projects()

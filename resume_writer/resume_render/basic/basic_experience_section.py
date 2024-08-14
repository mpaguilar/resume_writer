import logging
from datetime import datetime

import docx.document
from resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)

log = logging.getLogger(__name__)


class BasicRenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> list[str]:
        """Render role skills section."""

        log.debug("Rendering role skills.")
        _paragraph_lines = []
        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _skills_str = ", ".join(self.role.skills)
            _paragraph_lines.append(f"Skills: {_skills_str}")

        return _paragraph_lines

    def _details(self) -> list[str]:
        """Render role details section."""
        log.debug("Rendering role details.")
        _paragraph_lines = []
        # job category
        if self.role.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {self.role.job_category}")

        if self.role.location and self.settings.location:
            _paragraph_lines.append(f"Location: {self.role.location}")

        if self.role.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"Agency: {self.role.agency_name}")

        if self.role.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {self.role.employment_type}")

        return _paragraph_lines

    def _dates(self) -> list[str]:
        """Render role dates section."""

        log.debug("Rendering role dates.")
        _paragraph_lines = []

        # Start date
        if not self.role.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(self.role.start_date, "%B %Y")
        _paragraph_lines.append(f"Start Date: {_value}")

        # End date
        if self.role.end_date:
            _value = datetime.strftime(self.role.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_value}")

        return _paragraph_lines

    def render(self) -> None:
        """Render role overview/basics section."""

        _paragraph_lines = []

        log.debug("Rendering roles section.")
        # company name is required
        if not self.role.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Company: {self.role.company}")

        if not self.role.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Title: {self.role.title}")

        _detail_lines = self._details()
        if len(_detail_lines) > 0:
            _paragraph_lines.extend(_detail_lines)

        _date_lines = self._dates()
        if len(_date_lines) > 0:
            _paragraph_lines.extend(_date_lines)

        _skills_lines = self._skills()
        if len(_skills_lines) > 0:
            _paragraph_lines.extend(_skills_lines)

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class BasicRenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section."""
        log.debug("Rendering roles section.")
        for _role in self.roles:
            self.document.add_heading("Work History", level=2)
            _role_render = BasicRenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            )


class BasicRenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section."""

    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object."""
        log.debug("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> list[str]:
        """Render project overview section."""

        log.debug("Rendering project overview.")

        _paragraph_lines = []
        _overview = self.project.overview

        _url_line = ""
        if self.settings.url_description and _overview.url_description:
            _url_line = f"{_overview.url_description} "
        if self.settings.url and _overview.url:
            _url_line += f" ({_overview.url})"

        _url_line = _url_line.strip()
        if _url_line:
            _paragraph_lines.append(_url_line)

        if self.settings.start_date and _overview.start_date:
            _start_date = datetime.strftime(_overview.start_date, "%B %Y")
            _paragraph_lines.append(f"Start Date: {_start_date}")

        if self.settings.end_date and _overview.end_date:
            _end_date = datetime.strftime(_overview.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_end_date}")

        return _paragraph_lines

    def _skills(self) -> list[str]:
        """Render project skills section."""

        log.debug("Rendering project skills.")

        _paragraph_lines = []

        _skills = ", ".join(self.project.skills)
        _paragraph_lines.append(f"Skills: {_skills}")

        return _paragraph_lines

    def render(self) -> None:
        """Render project section."""

        log.debug("Rendering project section.")

        _paragraph_lines = []

        if self.settings.overview and self.project.overview:
            _paragraph_lines.extend(self._overview())

        if self.settings.description and self.project.description:
            _paragraph_lines.append(self.project.description.text)

        if self.settings.skills and len(self.project.skills) > 0:
            _paragraph_lines.extend(self._skills())

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class BasicRenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section."""

    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object."""

        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section."""

        log.debug("Rendering projects section.")
        for _project in self.projects:
            self.document.add_heading("Projects", level=2)
            _project_render = BasicRenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()


class BasicRenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object."""

        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section."""

        log.debug("Rendering experience section.")

        if self.settings.roles and self.experience.roles:
            BasicRenderRolesSection(
                document=self.document,
                roles=self.experience.roles,
                settings=self.settings.roles_settings,
            ).render()

        if self.settings.projects and self.experience.projects:
            BasicRenderProjectsSection(
                document=self.document,
                projects=self.experience.projects,
                settings=self.settings.projects_settings,
            ).render()

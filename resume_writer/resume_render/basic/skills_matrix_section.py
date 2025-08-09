import logging

import docx.document
from docx.enum.table import WD_ALIGN_VERTICAL
from utils.skills_matrix import SkillsMatrix

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.models.parsers import ParseContext
from resume_writer.resume_render.render_settings import (
    ResumeSkillsMatrixSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderSkillsMatrixBase,
)

log = logging.getLogger(__name__)


class RenderSkillsMatrixSection(ResumeRenderSkillsMatrixBase):
    """Render skills for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeSkillsMatrixSettings,
        parse_context: ParseContext,
    ):
        """Initialize skills render object.

        Args:
            document: The DOCX document to which the skills section will be added.
            experience: The experience data containing roles and skill usage history.
            settings: Configuration settings for rendering the skills matrix.
            parse_context: Contextual information used during parsing, not directly used here.

        Returns:
            None

        Notes:
            1. Validate that the provided parse_context is an instance of ParseContext.
            2. Store the parse_context for potential future use.
            3. Initialize the parent class (ResumeRenderSkillsMatrixBase) with the given parameters.

        """
        log.debug("Initializing functional skills render object.")
        assert isinstance(parse_context, ParseContext)
        self.parse_context = parse_context

        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render skills section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Check if the experience object has any roles; if not, raise a ValueError.
            2. Create a SkillsMatrix instance from the experience roles.
            3. If settings.all_skills is True, generate a matrix for all skills using the special value "*all*".
            4. Otherwise, generate a matrix only for the skills specified in settings.skills.
            5. Calculate the number of rows needed for the table based on the number of skills (each row holds two skills).
            6. Create a new table with the calculated number of rows and 4 columns, using the "Table Grid" style.
            7. Add a header row with bolded column titles: "Skill", "YOE (from - to)", "Skill", "YOE (from - to)".
            8. Iterate through the sorted skills and their data.
            9. For each skill, determine its row and column position in the table.
            10. Insert the skill name into the appropriate cell and center it vertically.
            11. Format the YOE string using the first_used and last_used dates, along with the calculated YOE value.
            12. Insert the formatted YOE string into the adjacent cell.
            13. Set the table to auto-fit its contents.

        """
        log.debug("Rendering functional skills section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a skills matrix.")

        _skills_matrix = SkillsMatrix(self.experience.roles)
        if self.settings.all_skills:
            _skills_matrix = _skills_matrix.matrix(["*all*"])
        else:
            _skills_matrix = _skills_matrix.matrix(self.settings.skills)

        # Calculate the number of rows needed for the table
        _num_rows = len(_skills_matrix) // 2 + len(_skills_matrix) % 2

        _table = self.document.add_table(rows=_num_rows + 1, cols=4, style="Table Grid")

        # Add header row
        _header_row = _table.rows[0]
        _header_row.cells[0].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[1].paragraphs[0].add_run("YOE (from - to)").bold = True
        _header_row.cells[2].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[3].paragraphs[0].add_run("YOE (from - to)").bold = True

        # Iterate through the sorted skills and their YOE values
        for i, (skill, data) in enumerate(_skills_matrix.items()):
            row_index = i // 2 + 1
            col_index = i % 2 * 2

            _table.cell(row_index, col_index).text = skill
            _table.cell(
                row_index,
                col_index,
            ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Calculate YOE and format the string

            first_used = (
                data["first_used"].strftime("%Y") if data["first_used"] else "N/A"
            )
            last_used = data["last_used"].strftime("%Y") if data["last_used"] else "N/A"
            yoe = f"{data['yoe']} ({first_used} - {last_used})"

            _table.cell(row_index, col_index + 1).text = yoe

        _table.autofit = True

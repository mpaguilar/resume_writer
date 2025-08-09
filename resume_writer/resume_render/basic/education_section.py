import logging

import docx.document
from docx.shared import Pt

from resume_writer.models.education import Degree, Education
from resume_writer.resume_render.render_settings import ResumeEducationSettings
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderDegreeBase,
    ResumeRenderEducationBase,
)

log = logging.getLogger(__name__)


class RenderDegreeSection(ResumeRenderDegreeBase):
    """Render a single degree section in a resume document.

    This class handles formatting and rendering of a single academic degree,
    including school, degree name, dates, major, and GPA, based on provided settings.

    Inherits from:
        ResumeRenderDegreeBase: Base class providing common rendering functionality.

    Attributes:
        document (docx.document.Document): The Word document being rendered.
        degree (Degree): The degree object containing education details.
        settings (ResumeEducationSettings): Configuration settings for rendering fields.
        font_size (int): Font size used for rendering (inherited from base class).
    """

    def __init__(
        self,
        document: docx.document.Document,
        degree: Degree,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic degree renderer.

        Args:
            document (docx.document.Document): The Word document to render the degree section into.
            degree (Degree): The degree object containing education details such as school, degree name, dates, major, and GPA.
            settings (ResumeEducationSettings): Configuration settings that control which degree fields are rendered (e.g., school, degree, dates, major, GPA).

        Returns:
            None

        Notes:
            1. Store the provided document, degree, and settings as instance attributes.
            2. No disk, network, or database access occurs during initialization.

        """
        super().__init__(document=document, degree=degree, settings=settings)

    def render(self) -> None:
        """Render a single degree with formatted details in the document.

        Args:
            None

        Returns:
            None

        Notes:
            1. Create a new paragraph in the document to hold the degree details.
            2. If the school name is present and enabled in settings, add it in bold, underlined, and larger font size.
            3. If the degree name is present and enabled in settings, add it in bold.
            4. If the start date is present and enabled in settings, format it as "Month Year" and add it.
            5. If the end date is present and enabled in settings, format it as "Month Year" (or "Present" if None) and add it.
            6. If the major is present and enabled in settings, add it.
            7. If the GPA is present and enabled in settings, add it in the format "GPA: X.XX".
            8. All text additions use runs and breaks to format the output properly.
            9. No disk, network, or database access occurs during rendering.

        """
        _paragraph = self.document.add_paragraph()

        if self.degree.school and self.settings.school:
            _school_run = _paragraph.add_run(f"{self.degree.school}")
            _school_run.bold = True
            _school_run.underline = True
            _school_run.font.size = Pt(self.font_size + 2)
            _school_run.add_break()

        if self.degree.degree and self.settings.degree:
            _degree_run = _paragraph.add_run(f"{self.degree.degree}")
            _degree_run.bold = True
            _degree_run.add_break()

        if self.degree.start_date and self.settings.start_date:
            _value = self.degree.start_date.strftime("%B %Y")
            _start_date_run = _paragraph.add_run(f"{_value}")
            if self.settings.end_date:
                _paragraph.add_run(" - ")
            else:
                _start_date_run.add_break()

        if self.degree.end_date and self.settings.end_date:
            if self.degree.end_date is None:
                _value = "Present"
            else:
                _value = self.degree.end_date.strftime("%B %Y")
            _end_date_run = _paragraph.add_run(f"{_value}")
            _end_date_run.add_break()

        if self.degree.major and self.settings.major:
            _degree_run = _paragraph.add_run(f"{self.degree.major}")
            _degree_run.add_break()

        if self.degree.gpa and self.settings.gpa:
            _gpa_run = _paragraph.add_run(f"GPA: {self.degree.gpa}")
            _gpa_run.add_break()


class RenderEducationSection(ResumeRenderEducationBase):
    """Render the entire education section of a resume document.

    This class manages the rendering of a resume's education section, including
    a heading and a list of degrees, with formatting based on user-defined settings.

    Inherits from:
        ResumeRenderEducationBase: Base class providing common education rendering behavior.

    Attributes:
        document (docx.document.Document): The Word document being rendered.
        education (Education): The education object containing a list of degrees and related data.
        settings (ResumeEducationSettings): Configuration settings for rendering behavior.
        font_size (int): Font size used for rendering (inherited from base class).
    """

    def __init__(
        self,
        document: docx.document.Document,
        education: Education,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic education renderer.

        Args:
            document (docx.document.Document): The Word document to render the education section into.
            education (Education): The education object containing a list of degrees and related data.
            settings (ResumeEducationSettings): Configuration settings that control rendering behavior, including whether to render degrees and which fields to include.

        Returns:
            None

        Notes:
            1. Store the provided document, education, and settings as instance attributes.
            2. No disk, network, or database access occurs during initialization.

        """
        super().__init__(document, education, settings)

    def render(self) -> None:
        """Render the education section with heading and degrees.

        Args:
            None

        Returns:
            None

        Notes:
            1. If the 'degrees' setting is disabled, exit early without rendering.
            2. Add a level-2 heading titled "Education" to the document.
            3. For each degree in the education object, create a RenderDegreeSection instance and call its render method.
            4. Rendering each degree is delegated to the RenderDegreeSection class.
            5. No disk, network, or database access occurs during rendering.

        """
        log.debug("Rendering education section.")

        if not self.settings.degrees:
            return

        self.document.add_heading("Education", level=2)
        for _degree in self.education.degrees:
            RenderDegreeSection(
                document=self.document,
                degree=_degree,
                settings=self.settings,
            ).render()

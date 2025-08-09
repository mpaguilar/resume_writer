import logging

import docx.document

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.resume_render.render_settings import (
    ResumeExecutiveSummarySettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExecutiveSummaryBase,
)

log = logging.getLogger(__name__)


class RenderExecutiveSummarySection(ResumeRenderExecutiveSummaryBase):
    """Render experience for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExecutiveSummarySettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document: The Word document object to which the executive summary will be added.
            experience: The experience data containing roles and related information.
            settings: Configuration settings for rendering the executive summary section.

        Returns:
            None

        Notes:
            1. Initialize the parent class with the provided document, experience, and settings.
            2. Log a debug message indicating the initialization of the functional experience render object.

        """
        log.debug("Initializing functional experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Log a debug message indicating the start of rendering the functional experience section.
            2. Validate that the experience object contains at least one role; raise a ValueError if not.
            3. Collect all unique job categories from the roles in the experience.
            4. For each job category specified in the settings:
                a. Filter roles belonging to the current category.
                b. If no roles are found for the category, log a warning and skip to the next category.
                c. Add a heading for the category with level 4.
                d. For each role in the category:
                    i. If no summary is available for the role, log a warning and skip to the next role.
                    ii. If no company is available for the role, log a warning and skip to the next role.
                    iii. Create a new paragraph with the bullet style.
                    iv. Add the role summary as a run to the paragraph.
                    v. Add the company name in italic as a run to the paragraph, appended to the summary.

        """
        log.debug("Rendering functional experience section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        # collect all job categories
        _roles = self.experience.roles
        _job_categories = set()

        # collect roles for each job category
        for _role in _roles:
            if _role.basics.job_category:
                _job_categories.add(_role.basics.job_category)

        # render each job category with roles

        for _category in self.settings.categories:
            _category_roles = [
                _role for _role in _roles if _role.basics.job_category == _category
            ]
            if len(_category_roles) == 0:
                _msg = f"No roles available for category {_category}"
                log.warning(_msg)
                continue
            self.document.add_heading(_category, level=4)

            for _role in _category_roles:
                if not _role.summary.summary:
                    _msg = f"No summary available for {_role.basics.title}"
                    log.warning(_msg)
                    continue

                if not _role.basics.company:
                    _msg = f"No company available for {_role.basics.title}"
                    log.warning(_msg)
                    continue

                _paragraph = self.document.add_paragraph()

                _paragraph.style = "List Bullet"
                _paragraph.add_run(f"{_role.summary.summary}")
                _run = _paragraph.add_run(f" ({_role.basics.company})")
                _run.italic = True

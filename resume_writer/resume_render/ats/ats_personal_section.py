import logging

from docx import Document
from models.resume import Resume
from resume_render.docx_render_base import ResumeRenderPersonalBase
from resume_render.render_settings import ResumePersonalSettings

log = logging.getLogger(__name__)


class ATSPersonalSection(ResumeRenderPersonalBase):
    """Represent and render personal section."""

    def __init__(
        self,
        document: Document,
        resume: Resume,
        settings: ResumePersonalSettings,
    ):
        """Initialize the object."""

        super().__init__(document=document, resume=resume, settings=settings)

    def contact_info(self) -> None:
        """Add contact information to the document."""

        if not self.resume.personal.contact_info:
            log.info("No contact information provided. Skipping.")
            return

        if not self.settings.contact_info:
            log.info("Contact information disabled. Skipping.")
            return

        self.document.add_heading("Contact information", 1)

        _paragraph_lines = []

        _contact = self.resume.personal.contact_info

        if not _contact.name:
            raise ValueError("Name is required for contact information.")

        _paragraph_lines.append(f"Name: {_contact.name}")

        if _contact.email:
            _paragraph_lines.append(f"Email: {_contact.email}")

        if _contact.phone:
            _paragraph_lines.append(f"Phone: {_contact.phone}")

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def website_info(self) -> None:
        """Add website info to the document."""

        if not self.resume.personal.websites:
            log.info("No website information provided. Skipping.")
            return

        if not self.settings.websites:
            log.info("Website information disabled. Skipping.")
            return

        _websites = self.resume.personal.websites

        _paragraph_lines = []

        self.document.add_heading("Websites", 2)

        if _websites.linkedin:
            _paragraph_lines.append(f"LinkedIn: {_websites.linkedin}")

        if _websites.website:
            _paragraph_lines.append(f"Website: {_websites.website}")

        if _websites.github:
            _paragraph_lines.append(f"GitHub: {_websites.github}")

        if _websites.twitter:
            _paragraph_lines.append(f"Twitter: {_websites.twitter}")

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def visa_status(self) -> None:
        """Add visa status to the document."""

        if not self.resume.personal.visa_status:
            log.info("No visa status provided. Skipping.")
            return

        if not self.settings.visa_status:
            log.info("Visa status disabled. Skipping.")
            return

        _visa = self.resume.personal.visa_status

        _paragraph_lines = []

        self.document.add_heading("Visa status", 2)

        if _visa.require_sponsorship is None:
            # the field doesn't exist
            pass
        else:
            _value = "Yes" if _visa.require_sponsorship else "No"
            _paragraph_lines.append(f"Require sponsorship: {_value}")

        if _visa.work_authorization is None:
            # the field doesn't exist
            pass
        else:
            _value = "Yes" if _visa.work_authorization else "No"
            _paragraph_lines.append(f"Authorized to work in the US: {_value}")

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def banner(self) -> None:
        """Add a banner to the document."""
        if not self.resume.personal.banner:
            log.info("No banner provided. Skipping.")
            return

        if not self.settings.banner:
            log.info("Banner disabled. Skipping.")
            return

        _banner = self.resume.personal.banner
        self.document.add_heading(_banner.text, 3)

    def note(self) -> None:
        """Add a note to the document."""
        if not self.resume.personal.note:
            log.info("No note provided. Skipping.")
            return

        if not self.settings.note:
            log.info("Note disabled. Skipping.")
            return

        _note = self.resume.personal.note
        self.document.add_heading(_note.text, 3)

    def render(self) -> None:
        """Render the personal section."""
        self.contact_info()

        if not self.resume.personal.banner and not self.resume.personal.note:
            log.info("No banner or note provided. Skipping.")
        else:
            self.banner()
            self.note()

        self.website_info()
        self.visa_status()

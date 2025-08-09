import logging
from datetime import datetime

import docx.document
from docx.shared import Pt

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)
from resume_writer.resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)

log = logging.getLogger(__name__)


class RenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document: The Docx document object to render into.
            role: The role data to render.
            settings: The settings for rendering the role.

        """
        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> list[str]:
        """Render role skills section.

        Returns:
            A list of strings containing the formatted skills text.

        Notes:
            1. Initialize an empty list to store the output lines.
            2. Check if the role has skills and the settings allow rendering skills.
            3. If skills exist, join them into a comma-separated string.
            4. Append the formatted "Skills: <skills>" string to the output list.
            5. Return the list of formatted skill lines.

        """
        log.debug("Rendering role skills.")
        _paragraph_lines = []
        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _skills_str = ", ".join(self.role.skills)
            _paragraph_lines.append(f"Skills: {_skills_str}")

        return _paragraph_lines

    def _details(self) -> list[str]:
        """Render role details section.

        Returns:
            A list of strings containing the formatted detail lines.

        Notes:
            1. Initialize an empty list to store the output lines.
            2. Extract the basics information from the role.
            3. If job category is present and enabled in settings, append it to the output list.
            4. If location is present and enabled in settings, append it to the output list.
            5. If agency name is present and enabled in settings, append it to the output list.
            6. If employment type is present and enabled in settings, append it to the output list.
            7. Return the list of formatted detail lines.

        """
        log.debug("Rendering role details.")
        _paragraph_lines = []
        # job category
        _basics = self.role.basics
        if _basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"{_basics.job_category}")

        if _basics.location and self.settings.location:
            _paragraph_lines.append(f"{_basics.location}")

        if _basics.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"{_basics.agency_name}")

        if _basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"{_basics.employment_type}")

        return _paragraph_lines

    def _dates(self) -> str:
        """Generate dates string for role.

        Returns:
            A formatted string containing start date and optional end date.

        Notes:
            1. Extract the basics information from the role.
            2. If start date is missing, log a warning and append an error.
            3. Format the start date as "mm-YYYY".
            4. If end date is present, format it as "mm-YYYY" and append to the string.
            5. If end date is missing, append " - Present" to the string.
            6. Return the final formatted date string.

        """
        log.debug("Rendering role dates.")

        _basics = self.role.basics
        # Start date
        if not _basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(_basics.start_date, "%m-%Y")
        _start_and_end = f"{_value}"

        # End date
        if _basics.end_date:
            _value = datetime.strftime(_basics.end_date, "%m-%Y")
            _start_and_end += f" - {_value}"
        else:
            _start_and_end += " - Present"

        return _start_and_end

    def render(self) -> None:
        """Render role overview/basics section.

        Notes:
            1. Initialize an empty list to store paragraph lines.
            2. Create a paragraph for the role basics.
            3. If company name is missing, log an error and append to errors list.
            4. Add the company name as bold and underlined text with increased font size.
            5. Insert a line break.
            6. If title is missing, log an error and append to errors list.
            7. Add the title as bold text.
            8. Insert a line break.
            9. Generate and add the formatted dates to the paragraph.
            10. Create a paragraph for the details section.
            11. Collect and clean the detail lines.
            12. If any detail lines exist, add them to the details paragraph.
            13. If summary is present and enabled, add it as italicized text with spacing.
            14. If responsibilities are present and enabled, add them with spacing.
            15. If skills are present and enabled, add them with spacing.

        """
        _paragraph_lines = []

        _basics_paragraph = self.document.add_paragraph()

        log.debug("Rendering roles section.")
        _basics = self.role.basics

        # Add the company name
        # company name is required
        if not _basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.error(_msg)

        _company_run = _basics_paragraph.add_run(f"{_basics.company}")
        _company_run.bold = True
        _company_run.underline = True
        _company_run.font.size = Pt(self.font_size + 2)

        # move to next line
        _basics_paragraph.add_run().add_break()

        # Add the title
        if not _basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.error(_msg)

        _title_run = _basics_paragraph.add_run(f"{_basics.title}")
        _title_run.bold = True

        # move to next line
        _basics_paragraph.add_run().add_break()

        # add the dates
        _date_line = self._dates()
        _basics_paragraph.add_run(f"{_date_line}")

        # add the details section such as job category, agency, etc.
        _details_paragraph = self.document.add_paragraph()

        _detail_lines = self._details()
        if len(_detail_lines) > 0:
            _paragraph_lines.extend(_detail_lines)

        _clean_lines = [_line.replace("\n\n", "\n") for _line in _paragraph_lines]

        if len(_paragraph_lines) > 0:
            _details_paragraph.add_run("\n".join(_clean_lines))

        # add the summary section
        if self.role.summary and self.settings.summary:
            _summary_paragraph = self.document.add_paragraph()
            _summary_paragraph.paragraph_format.space_after = Pt(self.font_size / 2)
            _summary_paragraph.paragraph_format.space_before = Pt(self.font_size / 2)
            _summary_run = _summary_paragraph.add_run(self.role.summary.summary)
            _summary_run.italic = True

        # add the responsibilities section
        if self.role.responsibilities and self.settings.responsibilities:
            _responsibilities_text = self.role.responsibilities.text.replace(
                "\n\n",
                "\n",
            )
            _responsibilites_paragraph = self.document.add_paragraph()
            _responsibilities_run = _responsibilites_paragraph.add_run(
                _responsibilities_text,
            )
            _responsibilites_paragraph.paragraph_format.space_after = Pt(
                self.font_size / 2,
            )
            _responsibilites_paragraph.paragraph_format.space_before = Pt(
                self.font_size / 2,
            )

        _skills_lines = self._skills()
        if len(_skills_lines) > 0:
            _skills_paragraph = self.document.add_paragraph()
            _skills_paragraph.add_run("\n".join(_skills_lines))


class RenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object.

        Args:
            document: The Docx document object to render into.
            roles: The list of role data to render.
            settings: The settings for rendering the roles.

        """
        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section.

        Notes:
            1. If no roles are present, log info and return.
            2. Add a heading "Work History" with level 2.
            3. For each role:
            4. If it's the first role, add a horizontal line.
            5. Render the role using RenderRoleSection.
            6. If it's not the last role, add a horizontal line and a blank paragraph.

        """
        log.debug("Rendering roles section.")
        if len(self.roles) > 0:
            self.document.add_heading("Work History", level=2)
        else:
            log.info("No roles found")
            return

        for _role in self.roles:
            if _role == self.roles[0]:
                _p = self.document.add_paragraph()
                self.add_horizontal_line(paragraph=_p)

            RenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            ).render()

            # add two blank lines between roles
            if _role != self.roles[-1]:
                _p = self.document.add_paragraph()
                self.add_horizontal_line(paragraph=_p)
                self.document.add_paragraph()


class RenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section."""

    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object.

        Args:
            document: The Docx document object to render into.
            project: The project data to render.
            settings: The settings for rendering the project.

        """
        log.debug("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> list[str]:
        """Render project overview section.

        Returns:
            A list of strings containing the formatted overview lines.

        Notes:
            1. Initialize an empty list to store output lines.
            2. Extract the project overview data.
            3. If URL description and URL are present and enabled, create a formatted URL line.
            4. If start date is present and enabled, format it and append.
            5. If end date is present and enabled, format it and append.
            6. Return the list of formatted lines.

        """
        log.debug("Rendering project overview.")

        _paragraph_lines = []
        _overview = self.project.overview

        _url_line = ""
        if self.settings.url_description and _overview.url_description:
            _url_line = f"{_overview.url_description} "
        if self.settings.url and _overview.url:
            _url_line += f" ({_overview.url})"

        _url_line = _url_line.strip()
        if _url_line:
            _paragraph_lines.append(_url_line)

        if self.settings.start_date and _overview.start_date:
            _start_date = datetime.strftime(_overview.start_date, "%B %Y")
            _paragraph_lines.append(f"Start Date: {_start_date}")

        if self.settings.end_date and _overview.end_date:
            _end_date = datetime.strftime(_overview.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_end_date}")

        return _paragraph_lines

    def _skills(self) -> list[str]:
        """Render project skills section.

        Returns:
            A list of strings containing the formatted skills text.

        Notes:
            1. Initialize an empty list to store output lines.
            2. Join the project skills into a comma-separated string.
            3. Append the formatted "Skills: <skills>" string to the output list.
            4. Return the list of formatted skill lines.

        """
        log.debug("Rendering project skills.")

        _paragraph_lines = []

        _skills = ", ".join(self.project.skills)
        _paragraph_lines.append(f"Skills: {_skills}")

        return _paragraph_lines

    def render(self) -> None:
        """Render project section.

        Notes:
            1. Initialize an empty list to store paragraph lines.
            2. If overview is present and enabled, collect its lines.
            3. If description is present and enabled, add it to the lines.
            4. If skills are present and enabled, collect their lines.
            5. If any lines exist, add them to the document as a single paragraph.

        """
        log.debug("Rendering project section.")

        _paragraph_lines = []

        if self.settings.overview and self.project.overview:
            _paragraph_lines.extend(self._overview())

        if self.settings.description and self.project.description:
            _paragraph_lines.append(self.project.description.text)

        if self.settings.skills and len(self.project.skills) > 0:
            _paragraph_lines.extend(self._skills())

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class RenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section."""

    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object.

        Args:
            document: The Docx document object to render into.
            projects: The list of project data to render.
            settings: The settings for rendering the projects.

        """
        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section.

        Notes:
            1. If no projects are present, log info and return.
            2. Add a heading "Projects" with level 2.
            3. For each project:
            4. Create a RenderProjectSection object.
            5. Render the project.

        """
        log.debug("Rendering projects section.")
        if len(self.projects) > 0:
            self.document.add_heading("Projects", level=2)
        for _project in self.projects:
            _project_render = RenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()


class RenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object.

        Args:
            document: The Docx document object to render into.
            experience: The experience data to render.
            settings: The settings for rendering the experience.

        """
        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section.

        Notes:
            1. If roles are present and enabled, render the roles section.
            2. If projects are present and enabled, render the projects section.

        """
        log.debug("Rendering experience section.")

        if self.settings.roles and self.experience.roles:
            RenderRolesSection(
                document=self.document,
                roles=self.experience.roles,
                settings=self.settings.roles_settings,
            ).render()

        if self.settings.projects and self.experience.projects:
            RenderProjectsSection(
                document=self.document,
                projects=self.experience.projects,
                settings=self.settings.projects_settings,
            ).render()

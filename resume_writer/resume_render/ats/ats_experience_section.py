import logging
from datetime import datetime

import docx.document
from models.experience import RoleBasics, RoleSkills
from models.resume import Resume
from resume_render.render_settings import ResumeRolesSettings
from resume_render.resume_render_base import RenderBase, ResumeRenderExperienceBase

log = logging.getLogger(__name__)


class ATSRenderExperienceSkillsSection(RenderBase):
    """Render skills section to a document."""

    def __init__(
        self,
        document: docx.document.Document,
        skills: RoleSkills,
    ):
        """Initialize the skills section."""

        assert isinstance(skills, RoleSkills)

        super().__init__()
        self.document = document
        self.skills = skills

    def render(self) -> None:
        """Render skills to the document."""
        log.debug("Rendering role skills")

        if len(self.skills) == 0:
            _msg = "No skills to render"
            self.warnings.append(_msg)
            log.info(_msg)
            return

        _skill_text = ", ".join(self.skills)

        self.document.add_paragraph(f"Skills: {_skill_text}")


class ATSRenderExperienceRoleBasicsSection(RenderBase):
    """Render role basics section to a document."""

    def __init__(
        self,
        document: docx.document.Document,
        basics: RoleBasics,
        settings: ResumeRolesSettings,
    ):
        """Initialize the role basics section."""

        assert isinstance(basics, RoleBasics)
        # Initialize the base class
        super().__init__()
        self.document = document
        self.basics = basics
        self.settings = settings

    def render(self) -> None: # noqa: C901
        """Render role basics to the document."""

        log.debug("Rendering role basics")

        _paragraph_lines = []

        # company name is required
        if not self.basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Company: {self.basics.company}")

        # Title is rendered by calling function
        if not self.basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Title: {self.basics.title}")

        # job category
        if self.basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {self.basics.job_category}")

        if self.basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {self.basics.location}")

        if self.basics.agency_name and self.settings.accomplishments:
            _paragraph_lines.append(f"Agency: {self.basics.agency_name}")

        if self.basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {self.basics.employment_type}")

        # Start date
        if not self.basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(self.basics.start_date, "%B %Y")
        _paragraph_lines.append(f"Start Date: {_value}")

        # End date
        if self.basics.end_date:
            _value = datetime.strftime(self.basics.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_value}")

        # Location
        if self.basics.location:
            _paragraph_lines.append(f"Location: {self.basics.location}")

        # Reason for change
        if self.basics.reason_for_change and self.settings.reason_for_leaving:
            _paragraph_lines.append(
                f"Reason for change: {self.basics.reason_for_change}",
            )

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class ATSExperienceSection(ResumeRenderExperienceBase):
    """Render roles section to a document."""

    def __init__(
        self,
        document: docx.document.Document,
        resume: Resume,
        settings: ResumeRolesSettings,
    ):
        """Initialize the roles section."""

        super().__init__(document=document, resume=resume, settings=settings)

    def roles(self) -> None:
        """Render roles to the document."""

        log.debug("Rendering roles")
        self.document.add_heading("Experience", 1)

        for _role in self.resume.experience.roles:
            if not _role.basics.title:
                self.errors.append("Role title is required")

            self.document.add_heading(_role.basics.title, 2)
            _basics = ATSRenderExperienceRoleBasicsSection(
                document=self.document,
                basics=_role.basics,
                settings=self.settings,
            )
            _basics.render()

            if _role.summary and self.settings.summary:
                self.document.add_paragraph(_role.summary.summary)

            if _role.responsibilities and self.settings.responsibilities:
                self.document.add_paragraph(_role.responsibilities.text)

            if _role.skills and self.settings.skills:
                _skills = ATSRenderExperienceSkillsSection(self.document, _role.skills)
                _skills.render()

    def render(self) -> None:
        """Render the section to the document."""
        if len(self.resume.experience.roles) == 0:
            _msg = "No roles to render"
            log.info(_msg)
            self.warnings.append(_msg)
            return

        self.roles()

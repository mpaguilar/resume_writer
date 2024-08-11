import logging
from datetime import datetime

import docx.document
from models.resume import Resume
from resume_render.docx_render_base import ResumeRenderCertificationsBase
from resume_render.render_settings import ResumeCertificationsSettings

log = logging.getLogger(__name__)


class ATSCertificationsSection(ResumeRenderCertificationsBase):
    """Represent and render certifications section."""

    def __init__(
        self,
        document: docx.document.Document,
        resume: Resume,
        settings: ResumeCertificationsSettings,
    ):
        """Initialize certification renderer."""

        assert isinstance(document, docx.document.Document)
        assert isinstance(resume, Resume)
        assert isinstance(settings, ResumeCertificationsSettings)

        super().__init__(document=document, resume=resume, settings=settings)

    def certifications(self) -> None:
        """Render certifications section."""

        if len(self.resume.certifications) == 0:
            log.info("No certifications listed. Skipping.")
            return

        self.document.add_heading("Certifications", level=1)

        for _cert in self.resume.certifications:
            _paragraph_lines = []

            # certificate name is required
            if not _cert.name:
                raise ValueError("Certificate name is required")

            self.document.add_heading(f"Certificate: {_cert.name}", level=2)

            if _cert.issuer:
                _paragraph_lines.append(f"Issued by: {_cert.issuer}")

            if _cert.issued:
                _value = datetime.strftime(_cert.issued, "%B %Y")
                _paragraph_lines.append(f"Issued: {_value}")

            if _cert.expires and self.settings.include_expires:
                _value = datetime.strftime(_cert.expires, "%B %Y")
                _paragraph_lines.append(f"Expires: {_value}")

            if len(_paragraph_lines) > 0:
                self.document.add_paragraph("\n".join(_paragraph_lines))

    def render(self) -> None:
        """Render certifications section."""
        self.certifications()

import logging

import docx.document
from docx.enum.table import WD_ALIGN_VERTICAL

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.models.parsers import ParseContext
from resume_writer.resume_render.render_settings import (
    ResumeSkillsMatrixSettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderSkillsMatrixBase,
)
from resume_writer.utils.skills_matrix import SkillsMatrix

log = logging.getLogger(__name__)


class RenderSkillsMatrixSection(ResumeRenderSkillsMatrixBase):
    """Render skills for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeSkillsMatrixSettings,
        parse_context: ParseContext,
    ):
        """Initialize skills render object.

        Args:
            document (docx.document.Document): The Word document to render into.
            experience (Experience): The parsed experience data containing roles and skill history.
            settings (ResumeSkillsMatrixSettings): Configuration settings for rendering skills matrix.
            parse_context (ParseContext): Contextual information used during parsing, used to track state.

        Returns:
            None

        Notes:
            1. Validate that the provided parse_context is an instance of ParseContext.
            2. Store the parse_context for later use during rendering.
            3. Call the parent class constructor with the provided document, experience, and settings.

        """
        log.debug("Initializing functional skills render object.")
        assert isinstance(parse_context, ParseContext)
        self.parse_context = parse_context

        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render skills section for functional resume.

        Args:
            None

        Returns:
            None

        Notes:
            1. Check if the experience object contains any roles; if not, raise a ValueError.
            2. Create a SkillsMatrix instance from the experience roles.
            3. If settings.all_skills is True, generate a matrix containing all skills; otherwise, use only the specified skills from settings.skills.
            4. Determine the number of rows needed in the table based on the number of skills (with two skills per row).
            5. Add a table with the calculated number of rows and 4 columns, using the "Table Grid" style.
            6. Add a header row with bolded column labels: "Skill", "YOE (from - to)", "Skill", and "YOE (from - to)".
            7. Iterate through the sorted skills and their associated data.
            8. For each skill, place it in the appropriate cell (alternating between columns 0 and 2).
            9. Format the years of experience (YOE) string as "{yoe} ({first_used} - {last_used})", using "N/A" if dates are missing.
            10. Place the formatted YOE string in the corresponding cell to the right of the skill.
            11. Set vertical center alignment for both skill and YOE cells.
            12. Automatically adjust table column widths to fit content.

        """
        log.debug("Rendering functional skills section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a skills matrix.")

        _skills_matrix = SkillsMatrix(self.experience.roles)
        if self.settings.all_skills:
            _skills_matrix = _skills_matrix.matrix(["*all*"])
        else:
            _skills_matrix = _skills_matrix.matrix(self.settings.skills)

        # Calculate the number of rows needed for the table
        _num_rows = len(_skills_matrix) // 2 + len(_skills_matrix) % 2

        _table = self.document.add_table(rows=_num_rows + 1, cols=4, style="Table Grid")

        # Add header row
        _header_row = _table.rows[0]
        _header_row.cells[0].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[1].paragraphs[0].add_run("YOE (from - to)").bold = True
        _header_row.cells[2].paragraphs[0].add_run("Skill").bold = True
        _header_row.cells[3].paragraphs[0].add_run("YOE (from - to)").bold = True

        # Iterate through the sorted skills and their YOE values
        for i, (skill, data) in enumerate(_skills_matrix.items()):
            row_index = i // 2 + 1
            col_index = i % 2 * 2

            _table.cell(row_index, col_index).text = skill
            _table.cell(
                row_index,
                col_index,
            ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Calculate YOE and format the string

            first_used = (
                data["first_used"].strftime("%Y") if data["first_used"] else "N/A"
            )
            last_used = data["last_used"].strftime("%Y") if data["last_used"] else "N/A"
            yoe = f"{data['yoe']} ({first_used} - {last_used})"

            _table.cell(row_index, col_index + 1).text = yoe

        _table.autofit = True

import logging

import docx.document
from docx.shared import Pt

from resume_writer.models.personal import ContactInfo, Personal
from resume_writer.resume_render.render_settings import ResumePersonalSettings
from resume_writer.resume_render.resume_render_base import ResumeRenderPersonalBase

log = logging.getLogger(__name__)


class RenderPersonalSection(ResumeRenderPersonalBase):
    """Render personal contact info section.

    Attributes:
        document (docx.document.Document): The Docx document object to render into.
        personal (Personal): The Personal model containing personal information.
        settings (ResumePersonalSettings): The ResumePersonalSettings object controlling which sections to render.
    """

    def __init__(
        self,
        document: docx.document.Document,
        personal: Personal,
        settings: ResumePersonalSettings,
    ):
        """Initialize the personal section renderer.

        Args:
            document (docx.document.Document): The Docx document object to render into.
            personal (Personal): The Personal model containing personal information.
            settings (ResumePersonalSettings): The ResumePersonalSettings object controlling which sections to render.

        Returns:
            None.

        Notes:
            1. Logs the initialization of the personal basic render object.
            2. Calls the parent class constructor with the provided arguments.
        """
        log.debug("Initializing personal basic render object")
        super().__init__(document, personal, settings)

    def _contact_info(self) -> None:
        """Render the contact info section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Creates a new paragraph in the document.
            2. Extracts contact info from the personal model.
            3. If name is present and enabled in settings, adds the name as a bold, larger font run and adds a line break.
            4. If email is present and enabled in settings, adds the email as a run and adds a line break.
            5. If phone is present and enabled in settings, adds the phone as a run and adds a line break.
            6. If location is present and enabled in settings, adds the location as a run and adds a line break.
        """
        log.debug("Rendering contact info section")

        _paragraph = self.document.add_paragraph()

        _info: ContactInfo = self.personal.contact_info
        if _info.name and self.settings.name:
            _name_run = _paragraph.add_run(f"{_info.name}")
            _name_run.bold = True
            _name_run.font.size = Pt(self.font_size + 2)
            _name_run.add_break()

        if _info.email and self.settings.email:
            _email_run = _paragraph.add_run(f"{_info.email}")
            _email_run.add_break()

        if _info.phone and self.settings.phone:
            _phone_run = _paragraph.add_run(f"{_info.phone}")
            _phone_run.add_break()

        if _info.location and self.settings.location:
            _location_run = _paragraph.add_run(f"{_info.location}")
            _location_run.add_break()

    def _banner(self) -> None:
        """Render the banner section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Extracts the banner text from the personal model.
            2. If banner text is present, adds a level 3 heading "Banner".
            3. Adds the banner text as a paragraph.
        """
        log.debug("Rendering banner section")

        _banner = self.personal.banner
        if _banner.text:
            self.document.add_heading("Banner", level=3)
            self.document.add_paragraph(_banner.text)

    def _note(self) -> None:
        """Render the note section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Extracts the note text from the personal model.
            2. If note text is present, adds a level 3 heading "Note".
            3. Adds the note text as a paragraph.
        """
        log.debug("Rendering note section")

        _note = self.personal.note
        if _note.text:
            self.document.add_heading("Note", level=3)
            self.document.add_paragraph(_note.text)

    def _websites(self) -> None:
        """Render the websites section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Initializes an empty list to store website lines.
            2. Extracts websites from the personal model.
            3. If GitHub is present and enabled in settings, appends "GitHub: <url>" to the list.
            4. If LinkedIn is present and enabled in settings, appends "LinkedIn: <url>" to the list.
            5. If website is present and enabled in settings, appends "Website: <url>" to the list.
            6. If Twitter is present and enabled in settings, appends "Twitter: <url>" to the list.
            7. If the list is not empty, adds a level 3 heading "Websites".
            8. Joins the list with newlines and adds it as a paragraph.
        """
        log.debug("Rendering websites section")

        _paragraph_lines = []

        _websites = self.personal.websites

        if _websites.github and self.settings.github:
            _paragraph_lines.append(f"GitHub: {_websites.github}")

        if _websites.linkedin and self.settings.linkedin:
            _paragraph_lines.append(f"LinkedIn: {_websites.linkedin}")

        if _websites.website and self.settings.website:
            _paragraph_lines.append(f"Website: {_websites.website}")

        if _websites.twitter and self.settings.twitter:
            _paragraph_lines.append(f"Twitter: {_websites.twitter}")

        if len(_paragraph_lines) > 0:
            self.document.add_heading("Websites", level=3)
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def _visa_status(self) -> None:
        """Render the visa status section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. Initializes an empty list to store visa status lines.
            2. Extracts visa status from the personal model.
            3. If work authorization is present and enabled in settings, appends "Work Authorization: <value>" to the list.
            4. If require_sponsorship is not None and enabled in settings, appends "Require Sponsorship: Yes" or "No" based on the value.
            5. If the list is not empty, adds a level 3 heading "Visa Status".
            6. Joins the list with newlines and adds it as a paragraph.
        """
        _paragraph_lines = []

        _visa_status = self.personal.visa_status

        if _visa_status.work_authorization and self.settings.work_authorization:
            _paragraph_lines.append(
                f"Work Authorization: {_visa_status.work_authorization}",
            )

        if (
            _visa_status.require_sponsorship is not None
            and self.settings.require_sponsorship
        ):
            _value = "Yes" if _visa_status.require_sponsorship else "No"
            _paragraph_lines.append(f"Require Sponsorship: {_value}")

        if len(_paragraph_lines) > 0:
            self.document.add_heading("Visa Status", level=3)
            self.document.add_paragraph("\n".join(_paragraph_lines))

    def render(self) -> None:
        """Render the personal section.

        Args:
            None.

        Returns:
            None.

        Notes:
            1. If contact info exists and enabled in settings, calls _contact_info.
            2. If banner exists and enabled in settings, calls _banner.
            3. If note exists and enabled in settings, calls _note.
            4. If websites exist and enabled in settings, calls _websites.
            5. If visa status exists and enabled in settings, calls _visa_status.
        """
        if self.personal.contact_info and self.settings.contact_info:
            self._contact_info()

        if self.personal.banner and self.settings.banner:
            self._banner()

        if self.personal.note and self.settings.note:
            self._note()

        if self.personal.websites and self.settings.websites:
            self._websites()

        if self.personal.visa_status and self.settings.visa_status:
            self._visa_status()

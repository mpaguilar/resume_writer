import logging

import docx.document
from docx.shared import Pt

from resume_writer.models.education import Degree, Education
from resume_writer.resume_render.render_settings import ResumeEducationSettings
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderDegreeBase,
    ResumeRenderEducationBase,
)

log = logging.getLogger(__name__)


class RenderDegreeSection(ResumeRenderDegreeBase):
    """Render a single academic degree section in a resume document.

    This class is responsible for formatting and adding a single degree's details
    to a Word document, including school name, degree type, dates, major, and GPA
    based on the provided settings.

    Args:
        document (docx.document.Document): The Word document object to which the degree section will be added.
        degree (Degree): The Degree object containing the academic details (school, degree, dates, major, GPA).
        settings (ResumeEducationSettings): Configuration object specifying which fields to render.

    Returns:
        None: This method does not return a value.

    Notes:
        1. The method initializes the rendering process for a single degree.
        2. It adds a paragraph to the document.
        3. If the school name is present and the school setting is enabled, it adds the school name in bold, underlined, and slightly larger font.
        4. If the degree name is present and the degree setting is enabled, it adds the degree name in bold.
        5. If the start date is present and the start date setting is enabled, it formats and adds the start date (e.g., "January 2020").
        6. If the end date is present and the end date setting is enabled, it formats and adds the end date (e.g., "December 2024") or "Present" if no end date is provided.
        7. If the major is present and the major setting is enabled, it adds the major name.
        8. If the GPA is present and the GPA setting is enabled, it adds the GPA label and value.
        9. The method does not perform any disk, network, or database access.

    """

    def __init__(
        self,
        document: docx.document.Document,
        degree: Degree,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic degree renderer."""
        super().__init__(document=document, degree=degree, settings=settings)

    def render(self) -> None:
        """Render a single degree.

        This method constructs and adds a formatted paragraph to the document
        representing a single degree, including school, degree, dates, major, and GPA
        as specified by the settings.

        Args:
            None: This method does not take any arguments.

        Returns:
            None: This method does not return a value.

        Notes:
            1. Logs debugging information about the rendering process.
            2. Creates a new paragraph in the document.
            3. Adds the school name (if enabled and present) with bold, underline, and increased font size.
            4. Adds a line break after the school name.
            5. Adds the degree name (if enabled and present) in bold.
            6. Adds a line break after the degree name.
            7. Adds the start date (if enabled and present) in "Month Year" format.
            8. Adds a " - " separator if the end date will be rendered.
            9. Adds a line break after the start date if no end date is rendered.
            10. Adds the end date (if enabled and present) in "Month Year" format or "Present" if no end date is provided.
            11. Adds a line break after the end date.
            12. Adds the major (if enabled and present).
            13. Adds a line break after the major.
            14. Adds the GPA (if enabled and present) with the label "GPA: ".
            15. The method does not perform any disk, network, or database access.

        """
        log.debug("Rendering degree section.")

        _paragraph = self.document.add_paragraph()

        if self.degree.school and self.settings.school:
            _school_run = _paragraph.add_run(f"{self.degree.school}")
            _school_run.bold = True
            _school_run.underline = True
            _school_run.font.size = Pt(self.font_size + 2)
            _school_run.add_break()

        if self.degree.degree and self.settings.degree:
            _degree_run = _paragraph.add_run(f"{self.degree.degree}")
            _degree_run.bold = True
            _degree_run.add_break()

        if self.degree.start_date and self.settings.start_date:
            _value = self.degree.start_date.strftime("%B %Y")
            _start_date_run = _paragraph.add_run(f"{_value}")
            if self.settings.end_date:
                _paragraph.add_run(" - ")
            else:
                _start_date_run.add_break()

        if self.degree.end_date and self.settings.end_date:
            if self.degree.end_date is None:
                _value = "Present"
            else:
                _value = self.degree.end_date.strftime("%B %Y")
            _end_date_run = _paragraph.add_run(f"{_value}")
            _end_date_run.add_break()

        if self.degree.major and self.settings.major:
            _degree_run = _paragraph.add_run(f"{self.degree.major}")
            _degree_run.add_break()

        if self.degree.gpa and self.settings.gpa:
            _gpa_run = _paragraph.add_run(f"GPA: {self.degree.gpa}")


class RenderEducationSection(ResumeRenderEducationBase):
    """Render the Education section of a resume in a Word document.

    This class is responsible for formatting and adding an entire education section
    to a Word document, including a heading and a list of rendered degree sections
    based on the provided settings.

    Args:
        document (docx.document.Document): The Word document object to which the education section will be added.
        education (Education): The Education object containing a list of Degree objects.
        settings (ResumeEducationSettings): Configuration object specifying which fields to render.

    Returns:
        None: This method does not return a value.

    Notes:
        1. Logs debugging information about the rendering process.
        2. Checks if the degrees setting is disabled; if so, returns early without rendering.
        3. Adds a level-2 heading labeled "Education" to the document.
        4. Iterates over each degree in the education's degrees list.
        5. For each degree, creates and calls RenderDegreeSection to render that degree.
        6. Adds a blank paragraph between degrees to separate them visually, except after the last degree.
        7. The method does not perform any disk, network, or database access.

    """

    def __init__(
        self,
        document: docx.document.Document,
        education: Education,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic education renderer."""
        super().__init__(document, education, settings)

    def render(self) -> None:
        """Render the education section.

        This method adds a formatted heading and a list of degree sections to the document,
        using the provided education data and rendering settings.

        Args:
            None: This method does not take any arguments.

        Returns:
            None: This method does not return a value.

        Notes:
            1. Logs debugging information about the rendering process.
            2. Checks if the degrees setting is disabled; if so, returns early.
            3. Adds a heading "Education" at level 2 to the document.
            4. Iterates through each degree in the education's degrees list.
            5. For each degree, instantiates a RenderDegreeSection and calls its render method.
            6. Adds a blank paragraph after each degree except the last one to provide visual separation.
            7. The method does not perform any disk, network, or database access.

        """
        log.debug("Rendering education section.")

        if not self.settings.degrees:
            return

        self.document.add_heading("Education", level=2)
        for _ndx, _degree in enumerate(self.education.degrees):
            RenderDegreeSection(
                document=self.document,
                degree=_degree,
                settings=self.settings,
            ).render()

            if _ndx < len(self.education.degrees) - 1:
                self.document.add_paragraph()

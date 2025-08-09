import logging

import docx.document

from resume_writer.models.experience import (
    Experience,
)
from resume_writer.resume_render.render_settings import (
    ResumeExecutiveSummarySettings,
)
from resume_writer.resume_render.resume_render_base import (
    ResumeRenderExecutiveSummaryBase,
)

log = logging.getLogger(__name__)


class RenderExecutiveSummarySection(ResumeRenderExecutiveSummaryBase):
    """Render executive summary section for a functional resume.

    This class is responsible for rendering the executive summary section of a resume based on experience data.
    It processes roles and their categories, formatting them into a structured document with headings and bullet points.

    Inherits from:
        ResumeRenderExecutiveSummaryBase: Base class providing common rendering functionality.

    Attributes:
        document (docx.document.Document): The Word document object to render into.
        experience (Experience): The experience data containing roles and their details.
        settings (ResumeExecutiveSummarySettings): The rendering settings that control which categories are included and how they are formatted.
    """

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExecutiveSummarySettings,
    ) -> None:
        """Initialize the executive summary render object.

        Args:
            document (docx.document.Document): The Word document object to render into.
            experience (Experience): The experience data containing roles and their details.
            settings (ResumeExecutiveSummarySettings): The rendering settings that control which categories are included and how they are formatted.

        Returns:
            None

        Notes:
            1. Logs a debug message indicating initialization.
            2. Calls the parent class constructor to initialize base rendering functionality.
        """
        log.debug("Initializing functional experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render the executive summary section of the resume.

        This method generates the executive summary by processing experience data and formatting it into a Word document.

        Args:
            None

        Returns:
            None

        Notes:
            1. Logs a debug message indicating the start of rendering.
            2. Checks if the experience object contains any roles; if not, raises a ValueError.
            3. Collects all unique job categories from the roles in the experience data.
            4. Iterates over the job categories specified in the settings.
            5. For each category, filters roles that belong to that category.
            6. If no roles are found for a category, logs a warning and skips to the next category.
            7. Adds a heading for the current job category to the document with level 4.
            8. For each role in the category:
               a. Checks if a summary is available; if not, logs a warning and skips.
               b. Checks if a company name is available; if not, logs a warning and skips.
               c. Creates a new paragraph with a bullet point style.
               d. Adds the role summary as plain text.
               e. Adds the company name in italics, appended to the summary.
            9. No network, disk, or database access occurs during execution.
        """
        log.debug("Rendering functional experience section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        # collect all job categories
        _roles = self.experience.roles
        _job_categories = set()

        # collect roles for each job category
        for _role in _roles:
            if _role.basics.job_category:
                _job_categories.add(_role.basics.job_category)

        # render each job category with roles

        for _category in self.settings.categories:
            _category_roles = [
                _role for _role in _roles if _role.basics.job_category == _category
            ]
            if len(_category_roles) == 0:
                _msg = f"No roles available for category {_category}"
                log.warning(_msg)
                continue
            self.document.add_heading(_category, level=4)

            for _role in _category_roles:
                if not _role.summary.summary:
                    _msg = f"No summary available for {_role.basics.title}"
                    log.warning(_msg)
                    continue

                if not _role.basics.company:
                    _msg = f"No company available for {_role.basics.title}"
                    log.warning(_msg)
                    continue

                _paragraph = self.document.add_paragraph()

                _paragraph.style = "List Bullet"
                _paragraph.add_run(f"{_role.summary.summary}")
                _run = _paragraph.add_run(f" ({_role.basics.company})")
                _run.italic = True

import logging
from datetime import datetime

import docx.document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)
from resume_writer.resume_render.skills_matrix import (
    skills_experience,
)

log = logging.getLogger(__name__)


class RenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> list[str]:
        """Render role skills section."""

        log.debug("Rendering role skills.")
        _paragraph_lines = []
        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _skills_str = ", ".join(self.role.skills)
            _paragraph_lines.append(f"Skills: {_skills_str}")

        return _paragraph_lines

    def _details(self) -> list[str]:
        """Render role details section."""
        log.debug("Rendering role details.")
        _paragraph_lines = []
        # job category
        _basics = self.role.basics
        if _basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {_basics.job_category}")

        if _basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {_basics.location}")

        if _basics.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"Agency: {_basics.agency_name}")

        if _basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {_basics.employment_type}")

        return _paragraph_lines

    def _dates(self) -> list[str]:
        """Render role dates section."""

        log.debug("Rendering role dates.")
        _paragraph_lines = []

        _basics = self.role.basics
        # Start date
        if not _basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(_basics.start_date, "%B %Y")
        _paragraph_lines.append(f"Start Date: {_value}")

        # End date
        if _basics.end_date:
            _value = datetime.strftime(_basics.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_value}")

        return _paragraph_lines

    def render(self) -> None:
        """Render role overview/basics section."""

        _paragraph_lines = []

        log.debug("Rendering roles section.")
        _basics = self.role.basics
        # company name is required
        if not _basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Company: {_basics.company}")

        if not _basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Title: {_basics.title}")

        _detail_lines = self._details()
        if len(_detail_lines) > 0:
            _paragraph_lines.extend(_detail_lines)

        _date_lines = self._dates()
        if len(_date_lines) > 0:
            _paragraph_lines.extend(_date_lines)

        if self.role.summary and self.settings.summary:
            self.document.add_paragraph(self.role.summary.summary)

        if self.role.responsibilities and self.settings.responsibilities:
            self.document.add_paragraph(self.role.responsibilities.text)

        _skills_lines = self._skills()
        if len(_skills_lines) > 0:
            _paragraph_lines.extend(_skills_lines)

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class RenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section."""
        log.debug("Rendering roles section.")
        if len(self.roles) > 0:
            self.document.add_heading("Work History", level=2)
        else:
            log.info("No roles found")
            return

        for _role in self.roles:
            RenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            ).render()


class RenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section."""

    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object."""
        log.debug("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> list[str]:
        """Render project overview section."""

        log.debug("Rendering project overview.")

        _paragraph_lines = []
        _overview = self.project.overview

        _url_line = ""
        if self.settings.url_description and _overview.url_description:
            _url_line = f"{_overview.url_description} "
        if self.settings.url and _overview.url:
            _url_line += f" ({_overview.url})"

        _url_line = _url_line.strip()
        if _url_line:
            _paragraph_lines.append(_url_line)

        if self.settings.start_date and _overview.start_date:
            _start_date = datetime.strftime(_overview.start_date, "%B %Y")
            _paragraph_lines.append(f"Start Date: {_start_date}")

        if self.settings.end_date and _overview.end_date:
            _end_date = datetime.strftime(_overview.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_end_date}")

        return _paragraph_lines

    def _skills(self) -> list[str]:
        """Render project skills section."""

        log.debug("Rendering project skills.")

        _paragraph_lines = []

        _skills = ", ".join(self.project.skills)
        _paragraph_lines.append(f"Skills: {_skills}")

        return _paragraph_lines

    def render(self) -> None:
        """Render project section."""

        log.debug("Rendering project section.")

        _paragraph_lines = []

        if self.settings.overview and self.project.overview:
            _paragraph_lines.extend(self._overview())

        if self.settings.description and self.project.description:
            _paragraph_lines.append(self.project.description.text)

        if self.settings.skills and len(self.project.skills) > 0:
            _paragraph_lines.extend(self._skills())

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


class RenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section."""

    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object."""

        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section."""

        log.debug("Rendering projects section.")
        if len(self.projects) > 0:
            self.document.add_heading("Projects", level=2)
        for _project in self.projects:
            _project_render = RenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()


class RenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object."""

        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section."""

        log.debug("Rendering experience section.")

        if self.settings.roles and self.experience.roles:
            RenderRolesSection(
                document=self.document,
                roles=self.experience.roles,
                settings=self.settings.roles_settings,
            ).render()

        if self.settings.projects and self.experience.projects:
            RenderProjectsSection(
                document=self.document,
                projects=self.experience.projects,
                settings=self.settings.projects_settings,
            ).render()


class RenderExecutiveSummarySkillsSection(ResumeRenderExperienceBase):
    """Render skills for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ):
        """Initialize skills render object."""
        log.debug("Initializing functional skills render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def find_skill_date_range(self, skill: str) -> tuple[datetime, datetime]:
        """Find the earliest job with a skill."""

        _earliest = None
        # collect the start dates for each role with this skill
        _start_dates = [
            role.basics.start_date
            for role in self.experience.roles
            if skill in role.skills
        ]
        # collect the end dates for each role with this skill
        _end_dates = [
            role.basics.end_date
            for role in self.experience.roles
            if skill in role.skills
        ]
        # find the earliest start date
        if len(_start_dates) > 0:
            _earliest_start_date = min(_start_dates)
            _last_end_date = max(_end_dates)

        return _earliest_start_date, _last_end_date

    def render(self) -> None:
        """Render skills section for functional resume."""

        log.debug("Rendering functional skills section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        _roles = self.experience.roles

        # use only skills specified in the settings
        _settings_skills = self.settings.executive_summary_settings.skills

        # get a dict of all skills and yoe
        _all_skills_yoe = skills_experience(_roles)

        _skills_yoe = {}
        # filter out skills not in the settings
        for _setting_skill in _settings_skills:
            _skills_yoe[_setting_skill] = _all_skills_yoe.get(_setting_skill, 0)

        # sort skills by yoe
        _skills_yoe = dict(
            sorted(_skills_yoe.items(), key=lambda item: item[1], reverse=True),
        )

        _num_rows = len(_skills_yoe) / 2
        _num_rows = int(_num_rows) + 1 if _num_rows % 1 > 0 else int(_num_rows)

        _skills_yoe_items = list(_skills_yoe.items())

        _table = self.document.add_table(rows=_num_rows, cols=4, style="Table Grid")

        # TODO: This will not scale with the font as-is
        _row_size = Inches(0.2)

        for x in range(_num_rows):
            _table.rows[x].height = _row_size

            _cell1 = _table.cell(x, 0)
            _cell1.text = _skills_yoe_items[x][0]
            _cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            _cell2_text = str(_skills_yoe_items[x][1])
            _first_date, _last_date = self.find_skill_date_range(
                _skills_yoe_items[x][0],
            )

            _first_date_str = _first_date.strftime("%Y")
            _last_date_str = _last_date.strftime("%Y")
            _cell2_text = _cell2_text + f" / ({_first_date_str} - {_last_date_str})"

            _cell2 = _table.cell(x, 1)
            _cell2.text = str(_cell2_text)
            _cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for x in range(_num_rows, len(_skills_yoe_items)):
            # add skill name to first column
            _cell1 = _table.cell(x - _num_rows, 2)
            _cell1.text = _skills_yoe_items[x][0]
            _cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            _cell2_text = str(_skills_yoe_items[x][1])

            _first_date, _last_date = self.find_skill_date_range(
                _skills_yoe_items[x][0],
            )

            _first_date_str = _first_date.strftime("%Y")
            _last_date_str = _last_date.strftime("%Y")
            _cell2_text = _cell2_text + f"  / ({_first_date_str} - {_last_date_str})"

            _cell2 = _table.cell(x - _num_rows, 3)
            _cell2.text = _cell2_text
            _cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _table.autofit = True


class RenderExecutiveSummaryExperienceSection(ResumeRenderExperienceBase):
    """Render experience for a functional resume."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object."""
        log.debug("Initializing functional experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section for functional resume."""

        log.debug("Rendering functional experience section.")

        if not self.experience.roles:
            raise ValueError("Experience must have roles for a functional resume.")

        # collect all job categories
        _roles = self.experience.roles
        _job_categories = set()

        # collect roles for each job category
        for _role in _roles:
            if _role.basics.job_category:
                _job_categories.add(_role.basics.job_category)

        # render each job category with roles

        for _category in self.settings.executive_summary_settings.categories:
            _category_roles = [
                _role for _role in _roles if _role.basics.job_category == _category
            ]
            self.document.add_heading(_category, level=4)

            for _role in _category_roles:
                _paragraph = self.document.add_paragraph()

                _paragraph.style = "List Bullet"
                _paragraph.add_run(f"{_role.summary.summary}")
                _run = _paragraph.add_run(f" ({_role.basics.company})")
                _run.italic = True

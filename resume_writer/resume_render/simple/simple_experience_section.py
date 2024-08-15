import logging
from datetime import datetime

import docx.document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt
from resume_render.render_settings import (
    ResumeExperienceSettings,
    ResumeProjectsSettings,
    ResumeRolesSettings,
)
from resume_render.resume_render_base import (
    ResumeRenderExperienceBase,
    ResumeRenderProjectBase,
    ResumeRenderProjectsBase,
    ResumeRenderRoleBase,
    ResumeRenderRolesBase,
)

from resume_writer.models.experience import (
    Experience,
    Project,
    Projects,
    Role,
    Roles,
)

log = logging.getLogger(__name__)


class BasicRenderRoleSection(ResumeRenderRoleBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        role: Role,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        log.debug("Initializing roles render object.")
        super().__init__(document=document, role=role, settings=settings)

    def _skills(self) -> None:
        """Render role skills section."""

        log.debug("Rendering role skills.")
        _paragraph_lines = []
        if self.role.skills and self.settings.skills and len(self.role.skills) > 0:
            _skills = [_s.strip() for _s in self.role.skills]
            _skills_str = ", ".join(_skills)
            _paragraph_lines.append(f"Skills: {_skills_str}")

        if len(_paragraph_lines) > 0:
            _skills_paragraph = self.document.add_paragraph()
            _skills_run = _skills_paragraph.add_run("\n".join(_paragraph_lines))
            _skills_run.font.size = Pt(self.font_size - 2)
            _skills_run.font.italic = True

            _skills_paragraph.paragraph_format.space_after = Pt(10)

    def _details(self) -> list[str]:
        """Render role details section."""
        log.debug("Rendering role details.")
        _paragraph_lines = []
        # job category
        _basics = self.role.basics
        if _basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {_basics.job_category}")

        if _basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {_basics.location}")

        if _basics.agency_name and self.settings.agency_name:
            _paragraph_lines.append(f"Agency: {_basics.agency_name}")

        if _basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {_basics.employment_type}")

        return _paragraph_lines

    def _dates(self) -> str:
        """Render role dates section."""

        log.debug("Rendering role dates.")
        _date_str = ""

        _basics = self.role.basics
        # Start date
        if not _basics.start_date:
            _msg = "Start date is missing."
            self.errors.append(_msg)
            log.warning(_msg)

        if self.settings.start_date:
            _value = datetime.strftime(_basics.start_date, "%B %Y")
            _date_str = f"{_value}"

        # End date
        if _basics.end_date and self.settings.end_date:
            _value = datetime.strftime(_basics.end_date, "%B %Y")
            _date_str = f"{_date_str} - {_value}"

        return _date_str

    def _summary(self) -> None:
        """Render role summary section."""

        if self.role.summary and self.settings.summary:
            _summary_paragraph = self.document.add_paragraph()
            _summary_paragraph.paragraph_format.space_before = Pt(4)
            _summary_paragraph.paragraph_format.space_after = Pt(4)
            _run = _summary_paragraph.add_run(self.role.summary.summary)
            _run.font.size = Pt(self.font_size)
            _run.font.italic = True

    def _responsibilities(self) -> None:
        """Render role responsibilities section."""

        if self.role.responsibilities and self.settings.responsibilities:
            _paragraph_text = self.role.responsibilities.text
            _paragraph_text = _paragraph_text.replace("\n\n", "\n")
            _doc_paragraph = self.document.add_paragraph()
            _doc_paragraph.add_run(_paragraph_text)
            _doc_paragraph.paragraph_format.space_after = Pt(10)

    def render(self) -> None:
        """Render role overview/basics section."""

        log.debug("Rendering roles section.")

        _paragraph_lines = []

        _basics = self.role.basics

        # role details
        if not _basics.title:
            _msg = "Role title is missing. Skipping."
            self.errors.append(_msg)
            log.warning(_msg)
            return

        _p = self.document.add_paragraph()
        _run = _p.add_run(_basics.title)
        _run.font.size = Pt(self.font_size + 1)
        _run.font.bold = True
        _run.font.underline = True

        # company name is required
        if not _basics.company:
            _msg = "Role company name is missing. Skipping."
            self.errors.append(_msg)
            log.warning(_msg)
            return

        # calculate space
        _tstops = _p.paragraph_format.tab_stops
        _pw = self.document.sections[0].page_width
        _tstops.add_tab_stop(int(_pw - (_pw / 6)), WD_TAB_ALIGNMENT.RIGHT)

        _p.add_run("\t")

        _run = _p.add_run(f"{_basics.company}")
        _run.font.size = Pt(self.font_size + 1)
        _run.font.bold = True

        _date_line = self._dates()
        if len(_date_line) > 0:
            _date_p = self.document.add_paragraph()
            _run = _date_p.add_run(f"({_date_line})")
            _run.font.size = Pt(self.font_size - 2)
            _run.font.italic = True

        _detail_lines = self._details()
        if len(_detail_lines) > 0:
            _paragraph_lines.extend(_detail_lines)

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))

        self._summary()

        self._responsibilities()

        self._skills()


class BasicRenderRolesSection(ResumeRenderRolesBase):
    """Render experience roles section."""

    def __init__(
        self,
        document: docx.document.Document,
        roles: Roles,
        settings: ResumeRolesSettings,
    ):
        """Initialize roles render object."""

        super().__init__(document=document, roles=roles, settings=settings)

    def render(self) -> None:
        """Render roles section."""
        log.debug("Rendering roles section.")
        if len(self.roles) > 0:
            self.document.add_heading("Professional Experience", level=2)
        else:
            log.info("No roles found")
            return

        for _role in self.roles:
            BasicRenderRoleSection(
                document=self.document,
                role=_role,
                settings=self.settings,
            ).render()
            _hline = self.document.add_paragraph()
            self.add_horizontal_line(_hline, 2)
            _hline.paragraph_format.space_after = Pt(10)


class BasicRenderProjectSection(ResumeRenderProjectBase):
    """Render experience project section."""

    def __init__(
        self,
        document: docx.document.Document,
        project: Project,
        settings: ResumeProjectsSettings,
    ):
        """Initialize project render object."""
        log.info("Initializing project render object.")
        super().__init__(document=document, project=project, settings=settings)

    def _overview(self) -> list[str]:
        """Render project overview section."""

        log.debug("Rendering project overview.")

        _paragraph_lines = []
        _overview = self.project.overview

        _doc_paragraph = self.document.add_paragraph()

        if self.settings.title and _overview.title:
            _run = _doc_paragraph.add_run(_overview.title)
            _run.bold = True
            _run.font.size = Pt(self.font_size + 2)

        if self.settings.start_date and _overview.start_date:
            _value = datetime.strftime(_overview.start_date, "%Y")
            _run = _doc_paragraph.add_run(f" ({_value})")

        if self.settings.url and _overview.url:
            _run = _doc_paragraph.add_run(f" ({_overview.url})")

        return _paragraph_lines

    def _skills(self) -> None:
        """Render project skills section."""

        log.debug("Rendering project skills.")

        _paragraph_lines = []

        _skills = ", ".join(self.project.skills)
        _paragraph = self.document.add_paragraph()
        _run = _paragraph.add_run(f"Skills: {_skills} ")
        _run.font.size = Pt(self.font_size - 2)
        _run.italic = True

        return _paragraph_lines

    def render(self) -> None:
        """Render project section."""

        log.debug("Rendering project section.")

        _paragraph_lines = []

        if self.settings.overview and self.project.overview:
            _paragraph_lines.extend(self._overview())

        if self.settings.description and self.project.description:
            _paragraph_text = self.project.description.text
            _paragraph_text = _paragraph_text.replace("\n\n", "\n")
            _run = self.document.add_paragraph().add_run(_paragraph_text)

        if (
            self.project.skills
            and self.settings.skills
            and len(self.project.skills) > 0
        ):
            self._skills()

        if len(_paragraph_lines) > 0:
            _p = self.document.add_paragraph("\n".join(_paragraph_lines))


class BasicRenderProjectsSection(ResumeRenderProjectsBase):
    """Render experience projects section."""

    def __init__(
        self,
        document: docx.document.Document,
        projects: Projects,
        settings: ResumeProjectsSettings,
    ):
        """Initialize projects render object."""

        log.debug("Initializing projects render object.")

        super().__init__(document=document, projects=projects, settings=settings)

    def render(self) -> None:
        """Render projects section."""

        log.info("Rendering projects section.")

        if self.document.styles["Normal"].font.size:
            _font_size = self.document.styles["Normal"].font.size.pt
        else:
            raise ValueError("Normal style font size not set.")

        if len(self.projects) > 0:
            self.document.add_heading("Personal Projects", level=2)
        else:
            log.info("No projects to render.")
            return

        for _project in self.projects:
            BasicRenderProjectSection(
                document=self.document,
                project=_project,
                settings=self.settings,
            ).render()
            self.document.add_paragraph().paragraph_format.space_after = Pt(4)


class BasicRenderExperienceSection(ResumeRenderExperienceBase):
    """Render experience section."""

    def __init__(
        self,
        document: docx.document.Document,
        experience: Experience,
        settings: ResumeExperienceSettings,
    ) -> None:
        """Initialize experience render object."""

        log.debug("Initializing experience render object.")
        super().__init__(document=document, experience=experience, settings=settings)

    def render(self) -> None:
        """Render experience section."""

        log.debug("Rendering experience section.")

        if self.settings.roles and self.experience.roles:
            BasicRenderRolesSection(
                document=self.document,
                roles=self.experience.roles,
                settings=self.settings.roles_settings,
            ).render()

        if self.settings.projects and self.experience.projects:
            BasicRenderProjectsSection(
                document=self.document,
                projects=self.experience.projects,
                settings=self.settings.projects_settings,
            ).render()

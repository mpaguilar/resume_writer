import logging

import docx.document
from resume_render.render_settings import ResumeEducationSettings
from resume_render.resume_render_base import (
    ResumeRenderDegreeBase,
    ResumeRenderEducationBase,
)

from resume_writer.models.education import Degree, Education

log = logging.getLogger(__name__)


class BasicRenderDegreeSection(ResumeRenderDegreeBase):
    """Render Degree Section."""

    def __init__(
        self,
        document: docx.document.Document,
        degree: Degree,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic degree renderer."""
        super().__init__(document=document, degree=degree, settings=settings)

    def _first_line(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Render the first line of the education section in a resume.

        Parameters
        ----------
        paragraph : docx.text.paragraph.Paragraph
            The paragraph object to add the education text to.

        Returns
        -------
        None

        Notes
        -----
        This function adds a new run to the provided paragraph and initializes
        an empty string for the education text.
        If a degree school is present and enabled, it is appended to the text.
        If a degree start date is present and enabled, it is appended to the text.
        If a degree end date is present and enabled, it is appended to the text.
        If no end date is present but enabled, "Present" or "Current"
        is appended to the text.
        Finally, the education text is added to the run and made bold.

        """
        _education_run = paragraph.add_run()

        _education_text = ""

        # first line
        if self.degree.school and self.settings.school:
            _education_text += f"{self.degree.school}"

        if self.degree.start_date and self.settings.start_date:
            _value = self.degree.start_date.strftime("%B %Y")
            if _education_text:
                _education_text += f" ({_value}"

        if self.degree.end_date and self.settings.end_date:
            _value = self.degree.end_date.strftime("%B %Y")
            if _education_text:
                _education_text += f" - {_value})"
            else:
                _education_text += f"({_value})"
        elif self.settings.end_date:
            if _education_text:
                _education_text += " - Present)"
            else:
                _education_text += "(Current)"

        _education_run.add_text(_education_text)
        _education_run.font.bold = True

    def _second_line(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Render the second line of the education section in a resume.

        Parameters
        ----------
        paragraph : docx.text.paragraph.Paragraph
            The paragraph to which the second line will be added.

        Returns
        -------
        None

        Notes
        -----
        This function initializes an empty string for the second line of the section.
        If a degree is present and enabled in settings,
        it appends it to the second line.
        If a major is present and enabled in settings, it appends it to the second line.
        If a GPA is present and enabled in settings, it appends it to the second line.
        If the second line is not empty, it adds it to the provided paragraph.

        """
        # TODO: why does '\r' work, and '\n' doesn't? (use run.add_break())
        _second_line = "\r"

        if self.degree.degree and self.settings.degree:
            _second_line += f"Degree: {self.degree.degree}"

        if self.degree.major and self.settings.major:
            _second_line += f", Major: {self.degree.major}"

        if self.degree.gpa and self.settings.gpa:
            _second_line += f", GPA: {self.degree.gpa}"

        if _second_line != "\r":
            _second_line_run = paragraph.add_run()
            _second_line_run.add_text(_second_line)

    def render(self) -> None:
        """Render a single degree section in a resume.

        Parameters
        ----------
        None

        Returns
        -------
        None

        Notes
        -----
        1. Adds a new paragraph to the document.
        2. Renders the first line of the degree section.
        3. Renders the second line of the degree section.

        """

        log.debug("Rendering degree section.")

        _education_paragraph = self.document.add_paragraph()

        # first line
        self._first_line(_education_paragraph)

        # second line
        self._second_line(_education_paragraph)


class BasicRenderEducationSection(ResumeRenderEducationBase):
    """Render Education Section."""

    def __init__(
        self,
        document: docx.document.Document,
        education: Education,
        settings: ResumeEducationSettings,
    ):
        """Initialize the basic education renderer."""
        super().__init__(document, education, settings)

    def render(self) -> None:
        """Render the education section of a resume.

        1. Check if there are any degrees to render. If not, return.
        2. Add a heading for the education section.
        3. For each degree, create a BasicRenderDegreeSection object and render it.

        Parameters
        ----------
        None

        Returns
        -------
        None

        """

        log.debug("Rendering education section.")

        if not self.settings.degrees:
            return

        self.document.add_heading("Education", level=2)
        for _degree in self.education.degrees:
            BasicRenderDegreeSection(
                document=self.document,
                degree=_degree,
                settings=self.settings,
            ).render()

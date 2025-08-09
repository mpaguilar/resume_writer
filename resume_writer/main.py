import logging
from pathlib import Path

import click
import docx
import rich
import tomli

from resume_writer.models.parsers import ParseContext
from resume_writer.models.personal import Personal
from resume_writer.models.resume import Resume
from resume_writer.renderers.html_renderer import RenderResumeHtml
from resume_writer.renderers.markdown_renderer import RenderResumeMarkdown
from resume_writer.resume_render.ats.resume_main import (
    RenderResume as AtsRenderResume,
)
from resume_writer.resume_render.basic.resume_main import (
    RenderResume as BasicRenderResume,
)
from resume_writer.resume_render.plain.resume_main import (
    RenderResume as PlainRenderResume,
)
from resume_writer.resume_render.render_settings import ResumeRenderSettings
from resume_writer.utils.resume_stats import DateStats

logging.basicConfig(level=logging.DEBUG)

log = logging.getLogger(__name__)


def print_personal(personal: Personal) -> None:
    """Print the personal information.

    Args:
        personal (Personal): The personal information object to print.

    Returns:
        None

    Notes:
        1. Extracts the name, email, and phone number from the personal object.
        2. Prints each field using rich formatting with bold labels.
    """
    rich.print(f"[bold]Name:[/bold] {personal.contact_info.name}")
    rich.print(f"[bold]Email:[/bold] {personal.contact_info.email}")
    rich.print(f"[bold]Phone:[/bold] {personal.contact_info.phone}")


def career_years_of_experience(resume: Resume) -> None:
    """Calculate and print the total years of experience from resume roles.

    Args:
        resume (Resume): The resume object containing role information.

    Returns:
        None

    Notes:
        1. Retrieves all roles from the resume.
        2. Initializes a DateStats object to manage date ranges.
        3. Iterates over each role and adds its start and end date range to the DateStats.
        4. Calculates the total years of experience using the date statistics.
        5. Prints the years of experience using rich formatting with a bold label.
    """
    _roles = resume.roles
    _date_stats = DateStats()
    for role in _roles:
        _date_stats.add_date_range(role.basics.start_date, role.basics.end_date)

    _yoe = _date_stats.years_of_experience

    rich.print(f"[bold]Years of Experience:[/bold] {_yoe}")


def dump_resume(resume: Resume) -> None:
    """Dump the entire resume data to the console.

    Args:
        resume (Resume): The resume object to dump.

    Returns:
        None

    Notes:
        1. Calls print_personal to display personal contact details.
        2. Calls career_years_of_experience to display total years of experience.
    """
    _personal = resume.personal
    print_personal(_personal)
    career_years_of_experience(resume)


def load_settings(settings_file: str) -> dict:
    """Load resume rendering settings from a TOML file.

    Args:
        settings_file (str): Path to the TOML settings file.

    Returns:
        dict: A dictionary containing the parsed settings.

    Notes:
        1. Converts the settings_file path to a Path object.
        2. Opens the TOML file in binary mode.
        3. Parses the TOML content using tomli.load.
        4. Prints the parsed settings using rich.
        5. Returns the settings dictionary.
        6. Disk access: Reads from the settings_file path.
    """
    _settings_file = Path(settings_file)

    with _settings_file.open("rb") as _f:
        _toml = tomli.load(_f)
        rich.print(_toml)
    return _toml


def basic_render(
    docx_doc: docx.document.Document,
    resume: Resume,
    settings: ResumeRenderSettings,
) -> None:
    """Render the resume using the basic rendering style.

    Args:
        docx_doc (docx.document.Document): The Word document object to render into.
        resume (Resume): The resume object containing the content to render.
        settings (ResumeRenderSettings): The rendering settings for the output.

    Returns:
        None

    Notes:
        1. Validates that all inputs are of the correct type.
        2. Logs the start of the rendering process.
        3. Creates a BasicRenderResume instance with the provided document, resume, and settings.
        4. Calls the render method on the renderer to generate the document.
        5. Logs the completion of the rendering process.
        6. Disk access: Saves the rendered document to the output file path.
    """
    assert isinstance(docx_doc, docx.document.Document)
    assert isinstance(resume, Resume)
    assert isinstance(settings, ResumeRenderSettings)

    log.info("Rendering simple resume")
    _renderer = BasicRenderResume(document=docx_doc, resume=resume, settings=settings)
    _renderer.render()
    log.info("Render of basic resume complete")


def ats_render(
    docx_doc: docx.document.Document,
    resume: Resume,
    settings: ResumeRenderSettings,
) -> None:
    """Render the resume using the ATS (Applicant Tracking System) rendering style.

    Args:
        docx_doc (docx.document.Document): The Word document object to render into.
        resume (Resume): The resume object containing the content to render.
        settings (ResumeRenderSettings): The rendering settings for the output.

    Returns:
        None

    Notes:
        1. Validates that all inputs are of the correct type.
        2. Logs the start of the rendering process.
        3. Creates an AtsRenderResume instance with the provided document, resume, and settings.
        4. Calls the render method on the renderer to generate the document.
        5. Logs the completion of the rendering process.
        6. Disk access: Saves the rendered document to the output file path.
    """
    assert isinstance(docx_doc, docx.document.Document)
    assert isinstance(resume, Resume)
    assert isinstance(settings, ResumeRenderSettings)

    log.info("Rendering simple resume")
    _renderer = AtsRenderResume(document=docx_doc, resume=resume, settings=settings)
    _renderer.render()
    log.info("Render of simple resume complete")


def plain_render(
    docx_doc: docx.document.Document,
    resume: Resume,
    settings: ResumeRenderSettings,
) -> None:
    """Render the resume using the plain (minimalist) rendering style.

    Args:
        docx_doc (docx.document.Document): The Word document object to render into.
        resume (Resume): The resume object containing the content to render.
        settings (ResumeRenderSettings): The rendering settings for the output.

    Returns:
        None

    Notes:
        1. Validates that all inputs are of the correct type.
        2. Logs the start of the rendering process.
        3. Creates a PlainRenderResume instance with the provided document, resume, and settings.
        4. Calls the render method on the renderer to generate the document.
        5. Logs the completion of the rendering process.
        6. Disk access: Saves the rendered document to the output file path.
    """
    assert isinstance(docx_doc, docx.document.Document)
    assert isinstance(resume, Resume)
    assert isinstance(settings, ResumeRenderSettings)

    log.info("Rendering plain resume")

    _renderer = PlainRenderResume(
        document=docx_doc,
        resume=resume,
        settings=settings,
    )
    _renderer.render()

    log.info("Render of plain resume complete")


def html_render(
    resume: Resume,
    settings: ResumeRenderSettings,
) -> None:
    """Render the resume as an HTML file.

    Args:
        resume (Resume): The resume object containing the content to render.
        settings (ResumeRenderSettings): The rendering settings for the output.

    Returns:
        None

    Notes:
        1. Validates that all inputs are of the correct type.
        2. Logs the start of the HTML rendering process.
        3. Creates a RenderResumeHtml instance with the resume and settings.
        4. Calls the render method to generate the HTML content.
        5. Saves the rendered HTML to a file at "data/html_resume.html".
        6. Logs the completion of the rendering process.
        7. Disk access: Writes to the file "data/html_resume.html".
    """
    assert isinstance(resume, Resume)
    assert isinstance(settings, ResumeRenderSettings)

    log.info("Rendering HTML resume")

    _html_renderer = RenderResumeHtml(
        resume=resume,
        settings=settings,
    )
    _html_renderer.render()
    _html_renderer.save(Path("data/html_resume.html"))

    log.info("Render of HTML resume complete.")


def markdown_render(
    resume: Resume,
    settings: ResumeRenderSettings,
) -> None:
    """Render the resume as a Markdown file.

    Args:
        resume (Resume): The resume object containing the content to render.
        settings (ResumeRenderSettings): The rendering settings for the output.

    Returns:
        None

    Notes:
        1. Validates that all inputs are of the correct type.
        2. Logs the start of the Markdown rendering process.
        3. Creates a RenderResumeMarkdown instance with the resume and settings.
        4. Calls the render method to generate the Markdown content.
        5. Saves the rendered Markdown to a file at "data/markdown_resume.md".
        6. Logs the completion of the rendering process.
        7. Disk access: Writes to the file "data/markdown_resume.md".
    """
    assert isinstance(resume, Resume)
    assert isinstance(settings, ResumeRenderSettings)

    log.info("Rendering Markdown resume")

    _markdown_renderer = RenderResumeMarkdown(
        resume=resume,
        settings=settings,
    )
    _markdown_renderer.render()
    _markdown_renderer.save(Path("data/markdown_resume.md"))

    log.info("Render of Markdown resume complete.")


def parse_text_resume(input_file: str) -> Resume:
    """Parse a text-based resume file and convert it into a Resume object.

    Args:
        input_file (str): Path to the text file containing the resume content.

    Returns:
        Resume: The parsed Resume object.

    Notes:
        1. Opens the input file and reads its content.
        2. Splits the content into lines while preserving line endings.
        3. Creates a ParseContext object with the lines and initial line number.
        4. Parses the resume content using the Resume.parse method.
        5. Returns the resulting Resume object.
        6. Disk access: Reads from the input_file path.
    """
    with open(input_file) as _f:
        _resume_text = _f.read()
        _resume_lines = _resume_text.splitlines(keepends=True)

    _parse_context = ParseContext(lines=_resume_lines, doc_line_num=0)

    _resume = Resume.parse(_parse_context)
    return _resume


@click.command()
@click.argument("input_file", type=click.Path(exists=True))
@click.option("--output-file", type=click.Path(), default="data/resume.docx")
@click.option(
    "--settings-file",
    type=click.Path(exists=True),
)
@click.option(
    "--resume-type",
    type=click.Choice(["ats", "basic", "plain", "html", "markdown"]),
    default="simple",
)
def main(
    input_file: str,
    output_file: str,
    settings_file: str,
    resume_type: str,
) -> None:
    """Convert a text resume to a .docx file with specified rendering style.

    Args:
        input_file (str): Path to the input text resume file.
        output_file (str): Path where the output .docx file will be saved.
        settings_file (str): Path to the TOML file containing rendering settings.
        resume_type (str): Type of resume to generate ("ats", "basic", "plain", "html", "markdown").

    Returns:
        None

    Notes:
        1. Loads the settings from the settings_file using load_settings.
        2. Creates a ResumeRenderSettings object and updates it from the loaded settings.
        3. Parses the input resume text file into a Resume object using parse_text_resume.
        4. Based on resume_type, selects the appropriate rendering method:
            a. "basic": Uses basic_render to generate a .docx file.
            b. "plain": Uses plain_render to generate a .docx file.
            c. "ats": Uses ats_render to generate a .docx file.
            d. "html": Uses html_render to generate an HTML file.
            e. "markdown": Uses markdown_render to generate a Markdown file.
        5. Saves the rendered output to the specified output_file path if applicable.
        6. Logs the completion of the process.
        7. Prints the entire resume object using rich.
        8. Disk access: Reads from input_file and settings_file; writes to output_file (for .docx) and data/html_resume.html, data/markdown_resume.md (for html/markdown).
    """
    _settings = load_settings(settings_file)
    _render_settings = ResumeRenderSettings()
    _render_settings.update_from_dict(_settings["resume"]["render"])

    _resume = parse_text_resume(input_file)

    if resume_type == "basic":
        _docx_doc = docx.Document()
        basic_render(_docx_doc, _resume, _render_settings)
        _docx_doc.save(output_file)
    elif resume_type == "plain":
        _docx_doc = docx.Document()
        plain_render(_docx_doc, _resume, _render_settings)
        _docx_doc.save(output_file)
    elif resume_type == "ats":
        _docx_doc = docx.Document()
        ats_render(_docx_doc, _resume, _render_settings)
        _docx_doc.save(output_file)
    elif resume_type == "html":
        html_render(_resume, _render_settings)
    elif resume_type == "markdown":
        markdown_render(resume=_resume, settings=_render_settings)
    else:
        raise ValueError(f"Unknown resume type: {resume_type}")

    log.info(f"Saved resume to {output_file}")

    rich.print(_resume)


if __name__ == "__main__":
    main()

import logging
from datetime import datetime

import docx.document
from docx_render.docx_render_base import DocxEducationBase
from docx_render.resume_settings import ResumeEducationSettings
from models.resume import Resume

log = logging.getLogger(__name__)


class ATSEducationSection(DocxEducationBase):
    """Renders the education section of a resume."""

    def __init__(
        self,
        document: docx.document.Document,
        resume: Resume,
        settings: ResumeEducationSettings,
    ):
        """Initialize the education rendering section."""
        assert isinstance(document, docx.document.Document)
        assert isinstance(resume, Resume)
        assert isinstance(settings, ResumeEducationSettings)

        super().__init__(document=document, resume=resume, settings=settings)

    def degrees(self) -> None:  # noqa: C901
        """Add Degrees section to document."""

        if not self.settings.degrees:
            log.info("Degrees section disabled. Skipping.")
            return

        if len(self.resume.education.degrees) == 0:
            log.info("No degrees found. Skipping.")
            return

        self.document.add_heading("Degrees", 2)

        for _degree in self.resume.education.degrees:
            # school name is required
            if not _degree.school:
                raise ValueError("School name is required")

            self.document.add_heading(f"School: {_degree.school}", 2)

            _paragraph_lines = []
            if _degree.degree:
                _paragraph_lines.append(f"Degree: {_degree.degree}")

            if _degree.start_date:
                _date_text = datetime.strftime(_degree.start_date, "%B %Y")
                _paragraph_lines.append(f"Start date: {_date_text}")

            if _degree.end_date:
                _date_text = datetime.strftime(_degree.end_date, "%B %Y")
                _paragraph_lines.append(f"End date: {_date_text}")

            if _degree.major:
                _paragraph_lines.append(f"Major: {_degree.major}")

            if _degree.gpa:
                _paragraph_lines.append(f"GPA: {_degree.gpa}")

            if len(_paragraph_lines) > 0:
                self.document.add_paragraph("\n".join(_paragraph_lines))

    def render(self) -> None:
        """Render the education section."""

        self.document.add_heading("Education", 1)
        if self.settings.degrees:
            self.degrees()

import logging
from datetime import datetime

import docx.document
from docx_render.docx_render_base import DocxRenderBase
from docx_render.resume_settings import ResumeRolesSettings
from models.resume import Resume
from models.roles import RoleBasics, RoleSkills

log = logging.getLogger(__name__)


class ATSRolesSection(DocxRenderBase):
    """Render roles section to a document."""

    def __init__(
        self,
        document: docx.document.Document,
        resume: Resume,
        settings: ResumeRolesSettings,
    ):
        """Initialize the roles section."""

        super().__init__()
        assert isinstance(resume, Resume)
        assert isinstance(document, docx.document.Document)
        assert isinstance(settings, ResumeRolesSettings)

        self.resume = resume
        self.document = document
        self.settings = settings

    def basics(self, basics: RoleBasics) -> None:  # noqa: C901
        """Render role basics to the document."""

        assert isinstance(basics, RoleBasics)

        log.debug("Rendering role basics")

        _paragraph_lines = []

        # company name is required
        if not basics.company:
            _msg = "Company name is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Company: {basics.company}")

        # Title is rendered by calling function
        if not basics.title:
            _msg = "Title is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _paragraph_lines.append(f"Title: {basics.title}")

        if basics.job_category and self.settings.job_category:
            _paragraph_lines.append(f"Job Category: {basics.job_category}")

        if basics.location and self.settings.location:
            _paragraph_lines.append(f"Location: {basics.location}")

        if basics.agency_name and self.settings.accomplishments:
            _paragraph_lines.append(f"Agency: {basics.agency_name}")

        if basics.employment_type and self.settings.employment_type:
            _paragraph_lines.append(f"Employment Type: {basics.employment_type}")

        # Start date
        if not basics.start_date:
            _msg = "Start date is required"
            self.errors.append(_msg)
            log.warning(_msg)

        _value = datetime.strftime(basics.start_date, "%B %Y")
        _paragraph_lines.append(f"Start Date: {_value}")

        # End date
        if basics.end_date:
            _value = datetime.strftime(basics.end_date, "%B %Y")
            _paragraph_lines.append(f"End Date: {_value}")

        if basics.location:
            _paragraph_lines.append(f"Location: {basics.location}")

        # Reason for change
        if basics.reason_for_change and self.settings.reason_for_leaving:
            _paragraph_lines.append(f"Reason for change: {basics.reason_for_change}")

        if len(_paragraph_lines) > 0:
            self.document.add_paragraph("\n".join(_paragraph_lines))


    def skills(self, skills: RoleSkills) -> None:
        """Render role skills to the document."""

        assert isinstance(skills, RoleSkills)

        log.debug("Rendering role skills")

        if len(skills) == 0:
            _msg = "No skills to render"
            self.warnings.append(_msg)
            log.info(_msg)
            return

        _skill_text = ", ".join(skills)

        self.document.add_paragraph(f"Skills: {_skill_text}")

    def roles(self) -> None:
        """Render roles to the document."""

        log.debug("Rendering roles")
        self.document.add_heading("Work History", 1)

        for _role in self.resume.roles:
            if not _role.basics.title:
                self.errors.append("Role title is required")

            self.document.add_heading(_role.basics.title, 2)
            self.basics(_role.basics)

            if _role.summary and self.settings.summary:
                self.document.add_paragraph(_role.summary.summary)

            if _role.responsibilities and self.settings.responsibilities:
                self.document.add_paragraph(_role.responsibilities.text)

            if _role.skills and self.settings.skills:
                self.skills(_role.skills)

    def render(self) -> None:
        """Render the section to the document."""
        if len(self.resume.roles) == 0:
            _msg = "No roles to render"
            log.info(_msg)
            self.warnings.append(_msg)
            return

        self.roles()

import logging

import docx.document
from docx_render.resume_settings import ResumeCertificationsSettings
from models.resume import Resume

log = logging.getLogger(__name__)


class ATSCertificationsSection:
    """Represent and render certifications section."""

    def __init__(
        self,
        document: docx.document.Document,
        resume: Resume,
        settings: ResumeCertificationsSettings,
    ):
        """Initialize certification renderer."""

        assert isinstance(document, docx.document.Document)
        assert isinstance(resume, Resume)
        assert isinstance(settings, ResumeCertificationsSettings)

        self.document = document
        self.settings = settings
        self.resume = resume

    def certifications(self) -> None:
        """Render certifications section."""

        if len(self.resume.certifications) == 0:
            log.info("No certifications listed. Skipping.")
            return

        self.document.add_heading("Certifications", level=1)

        for _cert in self.resume.certifications:
            _paragraph_lines = []

            # certificate name is required
            if not _cert.name:
                raise ValueError("Certificate name is required")

            self.document.add_heading(f"Certificate: {_cert.name}", level=2)

            if _cert.issuer:
                _paragraph_lines.append(f"Issued by: {_cert.issuer}")

            if _cert.issued:
                _paragraph_lines.append(f"Issued: {_cert.issued}")

            if _cert.expires and self.settings.include_expires:
                _paragraph_lines.append(f"Expires: {_cert.expires}")

            if len(_paragraph_lines) > 0:
                self.document.add_paragraph("\n".join(_paragraph_lines))

    def render(self) -> None:
        """Render certifications section."""

        self.certifications()

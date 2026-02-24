import re

import docx.text.paragraph


class MarkdownParser:
    """Convert markdown to OpenXML.

    Very limited, handles bold and italics, only.
    Limited to a single line, markdown spanning two lines will be ignored
    """

    @classmethod
    def parse_line(cls, 
                   paragraph : docx.text.paragraph.Paragraph,
                   line) -> docx.text.paragraph.Paragraph:

        bold_pattern = r"\*\*(.*?)\*\*"
        italic_pattern = r"\*(.*?)\*"

        # Replace bold and italic markers with unique tokens
        text = re.sub(bold_pattern, r"{{bold:\1}}", line)
        text = re.sub(italic_pattern, r"{{italic:\1}}", text)

        # Split the modified text into parts based on the tokens
        parts = re.split(r"({{bold:.*?}}|{{italic:.*?}})", text)

        parsed_list = []
        for part in parts:
            if part.startswith("{{bold:"):
                # Remove the token and record as bold
                _text_content = part[len("{{bold:") : -2]  # Strip '{{bold:' and '}}'
                _run = paragraph.add_run()
                _run.add_text(_text_content)
                _run.bold = True
                # parsed_list.append({"format": "bold", "text": text_content})
            elif part.startswith("{{italic:"):
                # Remove the token and record as italic
                _text_content = part[len("{{italic:") : -2]  # Strip '{{italic:' and '}}'
                _run = paragraph.add_run()
                _run.add_text(_text_content)
                _run.italic = True
                # parsed_list.append({"format": "italic", "text": text_content})
            elif part.strip():
                # Record as plain text if not empty
                _run = paragraph.add_run()
                _run.add_text(part)
                # parsed_list.append({"format": "plain", "text": part.strip()})

        paragraph.add_run().add_break()

        return paragraph
